#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Генератор отчетов
"""

from docx import Document
from docx.shared import Inches
from pathlib import Path
import sys
import os
from PIL import Image

# Добавляем корневую директорию в путь
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import table3
import table4
import table5
import prilog


class ReportGenerator:
    """Класс для генерации отчетов"""

    def __init__(self, session_manager):
        self.session_manager = session_manager
        self.project_root = Path(session_manager.project_root)
        self.templates_dir = self.project_root / "temp"

    def _replace_placeholders(self, doc):
        """Замена плейсхолдеров в документе на значения из ключевых элементов"""
        # Получаем данные ключевых элементов из БД
        settings_db_path = self.session_manager.get_settings_db_path()
        key_elements = {}

        try:
            import sqlite3
            conn = sqlite3.connect(settings_db_path)
            cursor = conn.cursor()
            cursor.execute("SELECT key, value FROM settings")
            rows = cursor.fetchall()
            conn.close()

            for key, value in rows:
                key_elements[key] = value or ''
        except Exception as e:
            print(f"Ошибка получения данных ключевых элементов: {e}")
            return

        # Маппинг плейсхолдеров к ключам в БД (с пробелами и с подчеркиваниями)
        placeholder_mapping = {
            # С пробелами
            '{{ НАИМЕНОВАНИЕ ОБЪЕКТА КАРТИРОВАНИЯ }}': key_elements.get('object_name', ''),
            '{{ НАИМЕНОВАНИЕ ОРГАНИЗАЦИИ ЗАЯВИТЕЛЯ }}': key_elements.get('organization_name', ''),
            '{{ ТЕМПЕРАТУРНЫЙ РЕЖИМ }}': key_elements.get('temp_mode', ''),
            '{{ ВЛАЖНОСТНЫЙ РЕЖИМ }}': key_elements.get('humidity_mode', ''),
            '{{ ДАТА ПРОВЕДЕНИЯ КАРТИРОВАНИЯ }}': key_elements.get('mapping_date', ''),
            '{{ ДАТА ВРЕМЯ ПРОВЕДЕНИЯ КАРТИРОВАНИЯ }}': key_elements.get('mapping_datetime', ''),
            '{{ ВИД КАРТИРОВАНИЯ }}': key_elements.get('mapping_type', ''),
            '{{ ДАТА ПОДПИСАНИЯ }}': key_elements.get('signature_date', ''),
            '{{ ДОЛЖНОСТЬ СОТРУДНИКА ФИРМЫ }}': key_elements.get('employee_position', ''),
            '{{ ФИО СОТРУДНИКА }}': key_elements.get('employee_name', ''),
            '{{ ПЛОЩАДЬ ПОМЕЩЕНИЯ }}': key_elements.get('area', ''),
            '{{ ВРЕМЯ ПРОВЕДЕНИЯ ИССЛЕДОВАНИЯ }}': key_elements.get('research_time', ''),
            '{{ НОМЕР ПРИЛОЖЕНИЯ СВИДЕТЕЛЬСТВА О ПОВЕРКЕ }}': key_elements.get('certificate_continuation_copy', ''),
            '{{ ОТЧЕТ ПО КАРТИРОВАНИЮ НАПИСАТЬ ПРОДОЛЖЕНИЕ }}': key_elements.get('certificate_continuation', ''),
            '{{ КОПИИ СВИДЕТЕЛЬСТВ О ПОВЕРКЕ СРЕДСТВ ИЗМЕРЕНИЙ И ПОВЫШЕНИЯ КВАЛИФИКАЦИИ СОТРУДНИКОВ ПРОДОЛЖИТЬ }}': key_elements.get('certificate_continuation_copy', ''),
            '{{ ДАТА ПРОВЕДЕНИЯ ПОВТОРОГО КАРТИРОВАНИЯ }}': key_elements.get('repeated_mapping_date', ''),
            '{{ ДАТА ПРОВЕДЕНИЯ ПОВТОРОНОГО КАРТИРОВАНИЯ }}': key_elements.get('repeated_mapping_date', ''),
            '{{ ИНТЕРВАЛ }}': key_elements.get('interval', ''),
            # С подчеркиваниями (альтернативный формат)
            '{{ НАИМЕНОВАНИЕ_ОБЪЕКТА_КАРТИРОВАНИЯ }}': key_elements.get('object_name', ''),
            '{{ НАИМЕНОВАНИЕ_ОРГАНИЗАЦИИ_ЗАЯВИТЕЛЯ }}': key_elements.get('organization_name', ''),
            '{{ ТЕМПЕРАТУРНЫЙ_РЕЖИМ }}': key_elements.get('temp_mode', ''),
            '{{ ВЛАЖНОСТНЫЙ_РЕЖИМ }}': key_elements.get('humidity_mode', ''),
            '{{ ДАТА_ПРОВЕДЕНИЯ_КАРТИРОВАНИЯ }}': key_elements.get('mapping_date', ''),
            '{{ ДАТА_ВРЕМЯ_ПРОВЕДЕНИЯ_КАРТИРОВАНИЯ }}': key_elements.get('mapping_datetime', ''),
            '{{ ВИД_КАРТИРОВАНИЯ }}': key_elements.get('mapping_type', ''),
            '{{ ДАТА_ПОДПИСАНИЯ }}': key_elements.get('signature_date', ''),
            '{{ ДОЛЖНОСТЬ_СОТРУДНИКА_ФИРМЫ }}': key_elements.get('employee_position', ''),
            '{{ ФИО_СОТРУДНИКА }}': key_elements.get('employee_name', ''),
            '{{ ПЛОЩАДЬ_ПОМЕЩЕНИЯ }}': key_elements.get('area', ''),
            '{{ ВРЕМЯ_ПРОВЕДЕНИЯ_ИССЛЕДОВАНИЯ }}': key_elements.get('research_time', ''),
            '{{ ВРЕ-МЯ_ПРОВЕДЕНИЯ_ИССЛЕДОВАНИЯ }}': key_elements.get('research_time', ''),
            '{{ НОМЕР_ПРИЛОЖЕНИЯ_СВИДЕТЕЛЬСТВА_О_ПОВЕРКЕ }}': key_elements.get('certificate_continuation_copy', ''),
            '{{ ОТЧЕТ_ПО_КАРТИРОВАНИЮ_НАПИСАТЬ_ПРОДОЛЖЕНИЕ }}': key_elements.get('certificate_continuation', ''),
            '{{ ОТЧЕТ_ПО_КАРТИРОВАНИЯ_НАПИСАТЬ_ПРОДОЛЖЕНИЕ }}': key_elements.get('certificate_continuation', ''),
            '{{ КОПИИ_СВИДЕТЕЛЬСТВ_О_ПОВЕРКЕ_СРЕДСТВ_ИЗМЕРЕНИЯ_И_ПОВЫШЕНИЯ_КВАЛИФИКАЦИИ_СОТРУДНИКОВ_ПРОДОЛЖИТЬ }}': key_elements.get('certificate_continuation_copy', ''),
            '{{ ДАТА_ПРОВЕДЕНИЯ_ПОВТОРОГО_КАРТИРОВАНИЯ }}': key_elements.get('repeated_mapping_date', ''),
            '{{ ДАТА_ПРОВЕДЕНИЯ_ПОВТОРОНОГО_КАРТИРОВАНИЯ }}': key_elements.get('repeated_mapping_date', ''),
            '{{ ИНТЕРВАЛ }}': key_elements.get('interval', ''),
        }
        
        # Все плейсхолдеры теперь будут иметь размер шрифта 12pt
        
        # Импортируем необходимые модули для работы со шрифтами
        from docx.shared import Pt

        # Функция для определения, находимся ли мы на первой странице
        def is_first_page(paragraph):
            """Определяем, находится ли параграф на первой странице"""
            # Надежная эвристика: считаем первые 15 параграфов первой страницей
            # Это покрывает типичный объем титульной страницы
            try:
                # Проверяем, является ли параграф первым в документе или одним из первых
                paragraph_index = doc.paragraphs.index(paragraph) if paragraph in doc.paragraphs else -1
                if paragraph_index >= 0 and paragraph_index < 15:  # Первые 15 параграфов считаем первой страницей
                    return True
                return False
            except:
                return False

        # Функция для определения, находится ли параграф в таблице
        def is_in_table(paragraph):
            """Определяем, находится ли параграф в таблице"""
            # Проверяем, является ли родительский элемент параграфа ячейкой таблицы
            try:
                parent = paragraph._p.getparent()
                if parent is not None:
                    # Проверяем, является ли родитель ячейкой таблицы
                    return parent.tag.endswith('tc')
                return False
            except:
                return False

        # Функция для применения размера шрифта 12pt ко всем плейсхолдерам
        def apply_font_size(run, placeholder, paragraph_context):
            """Устанавливаем размер шрифта 12pt для всех плейсхолдеров"""
            # Устанавливаем размер шрифта 12pt для всех плейсхолдеров независимо от контекста
            run.font.size = Pt(12)
            
            # Устанавливаем шрифт Times New Roman для всех случаев
            run.font.name = 'Times New Roman'

        # Получаем список фото
        photo_paths_str = key_elements.get('photo_paths', '')
        photo_list = [Path(p).name for p in photo_paths_str.split(',') if p.strip()] if photo_paths_str else []

        # Обрабатываем все параграфы в документе
        processed_paragraphs = set()
        bold_placeholders = [
            '{{ ОТЧЕТ ПО КАРТИРОВАНИЮ НАПИСАТЬ ПРОДОЛЖЕНИЕ }}',
            '{{ ДОЛЖНОСТЬ СОТРУДНИКА ФИРМЫ }}',
            '{{ ОТЧЕТ_ПО_КАРТИРОВАНИЮ_НАПИСАТЬ_ПРОДОЛЖЕНИЕ }}',
            '{{ ДОЛЖНОСТЬ_СОТРУДНИКА_ФИРМЫ }}',
        ]
        
        for paragraph in doc.paragraphs:
            if id(paragraph) in processed_paragraphs:
                continue
            
            # Определяем контекст параграфа
            paragraph_context = {
                'is_first_page': is_first_page(paragraph),
                'is_table': is_in_table(paragraph)
            }
                
            # Обрабатываем плейсхолдеры с жирным форматированием
            has_bold_placeholders = False
            for placeholder in bold_placeholders:
                if placeholder in paragraph.text and placeholder_mapping.get(placeholder, ''):
                    has_bold_placeholders = True
                    break
                    
            if has_bold_placeholders:
                # Для параграфов с жирными плейсхолдерами обрабатываем runs
                text = paragraph.text
                # Очищаем параграф
                paragraph.clear()
                # Позиция в тексте для обработки
                pos = 0
                while True:
                    # Найти ближайший плейсхолдер
                    min_pos = len(text)
                    found_placeholder = None
                    for placeholder in bold_placeholders:
                        p_pos = text.find(placeholder, pos)
                        if p_pos != -1 and p_pos < min_pos:
                            min_pos = p_pos
                            found_placeholder = placeholder
                    if not found_placeholder:
                        # Нет больше плейсхолдеров - добавить оставшийся текст
                        if pos < len(text):
                            paragraph.add_run(text[pos:])
                        break
                    # Добавить текст до плейсхолдера
                    if pos < min_pos:
                        paragraph.add_run(text[pos:min_pos])
                        # Добавить жирный плейсхолдер
                        value = placeholder_mapping[found_placeholder]
                        # Убираем лишний пробел в начале каждой строки для должности сотрудника и ФИО
                        if found_placeholder == '{{ ДОЛЖНОСТЬ СОТРУДНИКА ФИРМЫ }}' or found_placeholder == '{{ ДОЛЖНОСТЬ_СОТРУДНИКА_ФИРМЫ }}' or found_placeholder == '{{ ФИО СОТРУДНИКА }}' or found_placeholder == '{{ ФИО_СОТРУДНИКА }}':
                            if value:
                                lines = value.split('\n')
                                value = '\n'.join(line if line.strip() else line for line in lines)
                        run = paragraph.add_run(value)
                        run.bold = True
                        # Применяем контекстно-зависимый размер шрифта
                        apply_font_size(run, found_placeholder, paragraph_context)
                    # Добавляем отступ слева для текста из плейсхолдера
                    paragraph.paragraph_format.left_indent = Inches(0.1)  # Отступ слева
                    # Переместить позицию после текущего плейсхолдера
                    pos = min_pos + len(found_placeholder)
            else:
                # Для обычных параграфов используем простую замену
                for placeholder, value in placeholder_mapping.items():
                    if placeholder in paragraph.text:
                        # Убираем лишний пробел в начале каждой строки
                        if value:
                            lines = value.split('\n')
                            value_without_indent = '\n'.join(line if line.strip() else line for line in lines)
                        else:
                            value_without_indent = value
                        paragraph.text = paragraph.text.replace(placeholder, value_without_indent)
                        # Добавляем отступ слева для текста из плейсхолдера
                        if value and value.strip():
                            paragraph.paragraph_format.left_indent = Inches(0.1)  # Отступ слева
                        
            # Обрабатываем цикл фото
            import re
            photo_loop_pattern = r'\{\% for img in ФОТО \%\}([\s\S]*?)\{\% endfor \%\}'
            match = re.search(photo_loop_pattern, paragraph.text, re.IGNORECASE | re.DOTALL)
            if match:
                loop_content = match.group(1)
                if photo_list:
                    # Удаляем плейсхолдер из параграфа
                    paragraph.text = re.sub(photo_loop_pattern, '', paragraph.text, flags=re.IGNORECASE | re.DOTALL)
                    # Вставляем сами изображения в текущий параграф с выравниванием по центру
                    paragraph.alignment = 1  # 0 - левый, 1 - центрированный, 2 - правый
                    photos = [Path(p.strip()) for p in photo_paths_str.split(',') if p.strip()]
                    # Разбиваем на группы по 3 фото
                    for i in range(0, len(photos), 3):
                        row_photos = photos[i:i+3]
                        
                        for j, photo_path in enumerate(row_photos):
                            try:
                                run = paragraph.add_run()
                                img = Image.open(photo_path)
                                # Фиксированный размер: 4x5 см или 5x4 см в зависимости от ориентации
                                # 1 дюйм = 2.54 см
                                if img.width > img.height:
                                    # Горизонтальная фото: 3.5x2.8 см
                                    width_inches = 3.5 / 2.54
                                    height_inches = 2.8 / 2.54
                                else:
                                    # Вертикальная фото: 2.8x3.5 см
                                    width_inches = 2.8 / 2.54
                                    height_inches = 3.5 / 2.54
                                
                                # Масштабируем изображение до нужного размера с сохранением пропорций
                                img.thumbnail((int(width_inches * 96), int(height_inches * 96)), Image.Resampling.LANCZOS)
                                
                                # Сохраняем изображение во временную папку (если не существует)
                                temp_dir = Path("temp_photos")
                                temp_dir.mkdir(exist_ok=True)
                                temp_path = temp_dir / photo_path.name
                                img.save(temp_path)
                                
                                # Вставляем изображение в документ
                                run.add_picture(str(temp_path), width=Inches(width_inches), height=Inches(height_inches))
                                
                                # Добавляем отступ между фото (если не последняя в ряду)
                                if j < len(row_photos) - 1:
                                    run = paragraph.add_run()
                                    # Уменьшаем отступ между фото до минимального
                                    run.add_text(' ')  # Добавляем один пробел вместо табуляции
                            except Exception as e:
                                import traceback
                                traceback.print_exc()
                        
                        # Добавляем перевод строки между рядами
                        if i + 3 < len(photos):
                            paragraph.add_run().add_break()
                else:
                    # Если фото нет, просто удаляем плейсхолдер
                    paragraph.text = re.sub(photo_loop_pattern, '', paragraph.text, flags=re.IGNORECASE | re.DOTALL)
                    
                processed_paragraphs.add(id(paragraph))

        # Обрабатываем все таблицы в документе
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Обрабатываем плейсхолдеры с жирным форматированием в ячейках
                    has_bold_placeholders = False
                    for placeholder in bold_placeholders:
                        if placeholder in cell.text and placeholder_mapping.get(placeholder, ''):
                            has_bold_placeholders = True
                            break
                            
                    if has_bold_placeholders:
                        # Для ячеек с жирными плейсхолдерами обрабатываем runs в каждом параграфе
                        for paragraph in cell.paragraphs:
                            # Определяем контекст параграфа (в таблице)
                            paragraph_context = {
                                'is_first_page': is_first_page(paragraph),
                                'is_table': True  # Ячейка таблицы всегда в таблице
                            }
                            
                            text = paragraph.text
                            # Очищаем параграф
                            paragraph.clear()
                            # Позиция в тексте для обработки
                            pos = 0
                            while True:
                                # Найти ближайший плейсхолдер
                                min_pos = len(text)
                                found_placeholder = None
                                for placeholder in bold_placeholders:
                                    p_pos = text.find(placeholder, pos)
                                    if p_pos != -1 and p_pos < min_pos:
                                        min_pos = p_pos
                                        found_placeholder = placeholder
                                if not found_placeholder:
                                    # Нет больше плейсхолдеров - добавить оставшийся текст
                                    if pos < len(text):
                                        paragraph.add_run(text[pos:])
                                    break
                                # Добавить текст до плейсхолдера
                                if pos < min_pos:
                                    paragraph.add_run(text[pos:min_pos])
                                # Добавить жирный плейсхолдер
                                value = placeholder_mapping[found_placeholder]
                                run = paragraph.add_run(value)
                                run.bold = True
                                # Применяем контекстно-зависимый размер шрифта
                                apply_font_size(run, found_placeholder, paragraph_context)
                                # Добавляем отступ слева для текста из плейсхолдера
                                paragraph.paragraph_format.left_indent = Inches(0.1)  # Отступ слева
                                # Переместить позицию после текущего плейсхолдера
                                pos = min_pos + len(found_placeholder)
                    else:
                    # Для обычных ячеек используем простую замену
                        for placeholder, value in placeholder_mapping.items():
                            if placeholder in cell.text:
                                # Убираем лишний пробел в начале каждой строки
                                if value:
                                    lines = value.split('\n')
                                    value_without_indent = '\n'.join(line if line.strip() else line for line in lines)
                                else:
                                    value_without_indent = value
                                cell.text = cell.text.replace(placeholder, value_without_indent)
                                # Добавляем отступ слева для текста из плейсхолдера
                                for paragraph in cell.paragraphs:
                                    if value and value.strip():
                                        paragraph.paragraph_format.left_indent = Inches(0.1)  # Отступ слева
                    
                    # Центрируем текст в ячейках с определенными плейсхолдерами
                    center_placeholders = [
                        '{{ ДАТА ПОДПИСАНИЯ }}',
                        '{{ ДАТА_ПОДПИСАНИЯ }}',
                        '{{ ВРЕМЯ ПРОВЕДЕНИЯ ИССЛЕДОВАНИЯ }}',
                        '{{ ВРЕМЯ_ПРОВЕДЕНИЯ_ИССЛЕДОВАНИЯ }}',
                        '{{ ВРЕ-МЯ_ПРОВЕДЕНИЯ_ИССЛЕДОВАНИЯ }}',
                        '{{ ДАТА ПРОВЕДЕНИЯ КАРТИРОВАНИЯ }}',
                        '{{ ДАТА_ПРОВЕДЕНИЯ_КАРТИРОВАНИЯ }}'
                    ]
                    # Проверяем, содержит ли ячейка любой из нужных плейсхолдеров
                    has_center_placeholder = False
                    for placeholder in center_placeholders:
                        if placeholder in cell.text:
                            has_center_placeholder = True
                            break
                    
                    if has_center_placeholder:
                        from docx.enum.text import WD_ALIGN_PARAGRAPH
                        from docx.enum.table import WD_ALIGN_VERTICAL
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    if has_center_placeholder:
                        from docx.enum.text import WD_ALIGN_PARAGRAPH
                        from docx.enum.table import WD_ALIGN_VERTICAL
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                    # Обрабатываем цикл фото в ячейках таблиц
                    import re
                    photo_loop_pattern = r'\{\% for img in ФОТО \%\}([\s\S]*?)\{\% endfor \%\}'
                    match = re.search(photo_loop_pattern, cell.text, re.IGNORECASE | re.DOTALL)
                    if match:
                        loop_content = match.group(1)
                        if photo_list:
                            # Заменяем {{img}} на имена файлов
                            processed_content = loop_content
                            for i, photo_name in enumerate(photo_list):
                                processed_content = processed_content.replace('{{img}}', photo_name, 1)
                            # Если фото больше чем шаблонов, добавляем дополнительные
                            if len(photo_list) > processed_content.count(photo_name):
                                additional_content = ''
                                for photo in photo_list[len(processed_content.split(photo_name)) - 1:]:
                                    additional_content += loop_content.replace('{{img}}', photo)
                                processed_content += additional_content
                        else:
                            processed_content = ''
                        cell.text = re.sub(photo_loop_pattern, processed_content, cell.text, flags=re.IGNORECASE | re.DOTALL)

        # Обрабатываем вложенные таблицы и параграфы в ячейках
        def process_cell(cell):
            for paragraph in cell.paragraphs:
                # Обрабатываем плейсхолдеры с жирным форматированием в вложенных ячейках
                has_bold_placeholders = False
                for placeholder in bold_placeholders:
                    if placeholder in paragraph.text and placeholder_mapping.get(placeholder, ''):
                        has_bold_placeholders = True
                        break
                        
                if has_bold_placeholders:
                    # Определяем контекст параграфа (в таблице)
                    paragraph_context = {
                        'is_first_page': is_first_page(paragraph),
                        'is_table': True  # Вложенная ячейка таблицы всегда в таблице
                    }
                    
                    text = paragraph.text
                    # Очищаем параграф
                    paragraph.clear()
                    # Позиция в тексте для обработки
                    pos = 0
                    while True:
                        # Найти ближайший плейсхолдер
                        min_pos = len(text)
                        found_placeholder = None
                        for placeholder in bold_placeholders:
                            p_pos = text.find(placeholder, pos)
                            if p_pos != -1 and p_pos < min_pos:
                                min_pos = p_pos
                                found_placeholder = placeholder
                        if not found_placeholder:
                            # Нет больше плейсхолдеров - добавить оставшийся текст
                            if pos < len(text):
                                paragraph.add_run(text[pos:])
                            break
                        # Добавить текст до плейсхолдера
                        if pos < min_pos:
                            paragraph.add_run(text[pos:min_pos])
                        # Добавить жирный плейсхолдер
                        value = placeholder_mapping[found_placeholder]
                        run = paragraph.add_run(value)
                        run.bold = True
                        # Применяем контекстно-зависимый размер шрифта
                        apply_font_size(run, found_placeholder, paragraph_context)
                        # Добавляем отступ слева для текста из плейсхолдера
                        paragraph.paragraph_format.left_indent = Inches(0.1)  # Отступ слева
                        # Переместить позицию после текущего плейсхолдера
                        pos = min_pos + len(found_placeholder)
                else:
                    # Обрабатываем обычные плейсхолдеры
                    for placeholder, value in placeholder_mapping.items():
                        if placeholder in paragraph.text:
                            # Убираем лишний пробел в начале каждой строки
                            if value:
                                lines = value.split('\n')
                                value_without_indent = '\n'.join(line if line.strip() else line for line in lines)
                            else:
                                value_without_indent = value
                            paragraph.text = paragraph.text.replace(placeholder, value_without_indent)
                            # Добавляем отступ слева для текста из плейсхолдера
                            if value and value.strip():
                                paragraph.paragraph_format.left_indent = Inches(0.1)  # Отступ слева

                # Обрабатываем цикл фото в параграфах ячеек
                import re
                photo_loop_pattern = r'\{\% for img in ФОТО \%\}([\s\S]*?)\{\% endfor \%\}'
                match = re.search(photo_loop_pattern, paragraph.text, re.IGNORECASE | re.DOTALL)
                if match:
                    loop_content = match.group(1)
                    if photo_list:
                        # Заменяем {{img}} на имена файлов
                        processed_content = loop_content
                        for i, photo_name in enumerate(photo_list):
                            processed_content = processed_content.replace('{{img}}', photo_name, 1)
                        # Если фото больше чем шаблонов, добавляем дополнительные
                        if len(photo_list) > processed_content.count(photo_name):
                            additional_content = ''
                            for photo in photo_list[len(processed_content.split(photo_name)) - 1:]:
                                additional_content += loop_content.replace('{{img}}', photo)
                            processed_content += additional_content
                    else:
                        processed_content = ''
                    paragraph.text = re.sub(photo_loop_pattern, processed_content, paragraph.text, flags=re.IGNORECASE | re.DOTALL)

            for nested_table in cell.tables:
                for row in nested_table.rows:
                    for nested_cell in row.cells:
                        process_cell(nested_cell)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    process_cell(cell)
    
    def generate_report(self, report_type, template_path, output_path, 
                       use_humidity=False, other_info=None, periods=None):
        """
        Генерация отчета
        
        Args:
            report_type: Тип отчета ("Объект хранения", "Зона хранения", "Холодильник/Морозильник")
            template_path: Путь к шаблону
            output_path: Путь для сохранения отчета
            use_humidity: Учитывать влажность
            other_info: Дополнительная информация
            periods: Список периодов
        """
        try:
            # Загружаем шаблон
            doc = Document(template_path)

            # Заменяем плейсхолдеры на значения из ключевых элементов
            self._replace_placeholders(doc)
            
            # Определяем, какой модуль использовать для создания таблиц
            if report_type == "Объект хранения":
                table_module = table3
                table_func = table3.create_dynamic_tables3
            elif report_type == "Зона хранения":
                table_module = table4
                table_func = table4.create_dynamic_tables4
            elif report_type == "Холодильник/Морозильник":
                table_module = table5
                table_func = table5.create_dynamic_tables5
            else:
                raise ValueError(f"Неизвестный тип отчета: {report_type}")
            
            # Получаем данные из БД
            periods_db_path = self.session_manager.get_periods_db_path()
            logger_stats_db_path = self.session_manager.get_logger_stats_db_path()
            
            # Параметры из other_info
            if other_info:
                mapping_results = other_info.get('mapping_results', '')
                conclusion = other_info.get('conclusion', '')
                contract_date = other_info.get('contract_date', '')
                temp_homogeneity_text = other_info.get('temp_homogeneity_text')
                hum_homogeneity_text = other_info.get('hum_homogeneity_text')
                selected_recommendations = other_info.get('selected_recommendations', [])
            else:
                mapping_results = ''
                conclusion = ''
                contract_date = ''
                temp_homogeneity_text = None
                hum_homogeneity_text = None
                selected_recommendations = []
            
            # Диапазоны (получаем из БД)
            import sqlite3
            conn = sqlite3.connect(periods_db_path)
            cursor = conn.cursor()
            cursor.execute("SELECT DISTINCT required_mode_from, required_mode_to FROM periods LIMIT 1")
            temp_data = cursor.fetchone()
            conn.close()

            # Диапазоны температуры из БД
            temp_min = temp_data[0] if temp_data and temp_data[0] is not None else None
            temp_max = temp_data[1] if temp_data and temp_data[1] is not None else None

            # Диапазоны влажности: для Объекта/Зоны берём из other_info (таблица влажности всегда создаётся)
            humidity_min = None
            humidity_max = None
            if other_info and report_type in ("Объект хранения", "Зона хранения"):
                humidity_min = other_info.get('humidity_min')
                humidity_max = other_info.get('humidity_max')
            
            # Получаем выбранные периоды
            if periods:
                selected_periods = periods
            else:
                # Получаем все периоды из БД
                conn = sqlite3.connect(periods_db_path)
                cursor = conn.cursor()
                cursor.execute("SELECT id FROM periods")
                selected_periods = cursor.fetchall()
                conn.close()
            
            # Создаем таблицы
            if report_type == "Холодильник/Морозильник":
                table_func(
                    doc, selected_periods, periods_db_path, logger_stats_db_path,
                    temp_min, temp_max,
                    mapping_results, conclusion, contract_date,
                    temp_homogeneity_text, selected_recommendations
                )
            else:
                table_func(
                    doc, selected_periods, periods_db_path, logger_stats_db_path,
                    temp_min, temp_max, humidity_min, humidity_max,
                    mapping_results, conclusion, contract_date,
                    temp_homogeneity_text, hum_homogeneity_text, selected_recommendations
                )
            
            # Добавляем приложения
            images = {
                'layout': other_info.get('layout_image') if other_info else None,
                'loggers': other_info.get('loggers_image') if other_info else None,
                'temp_map': other_info.get('temp_map_image') if other_info else None,
                'humidity_map': other_info.get('humidity_map_image') if other_info else None
            }
            
            saved_risk_areas = other_info.get('risk_areas', []) if other_info else []
            
            # Создаем объект-заглушку с методом get()
            class TemplateVar:
                def __init__(self, value):
                    self._value = value
                def get(self):
                    return self._value
            
            selected_template_var = TemplateVar(report_type)
            period_images = {}  # TODO: добавить графики периодов
            use_humidity = other_info.get('use_humidity', False) if other_info else False
            logger_screenshots = other_info.get('logger_screenshots', []) if other_info else []
            
            # Получаем ориентацию изображений из other_info
            image_orientations = {}
            if other_info and 'image_orientations' in other_info:
                image_orientations = other_info['image_orientations']
            else:
                # Если ориентации не переданы, используем книжную по умолчанию
                image_orientations = {
                    'layout': 'portrait',
                    'loggers': 'portrait',
                    'temp_map': 'portrait',
                    'humidity_map': 'portrait'
                }

            prilog.create_appendices(
                doc, images, saved_risk_areas, selected_template_var, period_images,
                selected_recommendations, use_humidity=use_humidity,
                logger_screenshots=logger_screenshots,
                image_orientations=image_orientations
            )
            
            # Сохраняем отчет
            doc.save(output_path)
            
            return True
            
        except Exception as e:
            print(f"Ошибка генерации отчета: {e}")
            import traceback
            traceback.print_exc()
            return False
