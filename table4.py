import sqlite3
from docx.shared import Inches, Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
from docx.shared import Pt, Cm, Inches
from docx.shared import Mm

NEW_DIR = "new"


def left_align_cell_text(cell):
    """Align text to the left horizontally and center it vertically in a cell."""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
def set_cell_background(cell, color_hex):
    """Задать цвет фона ячейке"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def center_cell_text(cell):
    """Центрировать текст в ячейке"""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
def create_dynamic_tables4(doc, selected_periods, periods_db_path, logger_stats_db_path,
                           temp_min, temp_max, humidity_min, humidity_max,
                           mapping_results, conclusion, contract_date,
                           temp_homogeneity_text=None, hum_homogeneity_text=None,
                           selected_recommendations=None):
    # Создание сводной таблицы периодов (только один раз)
    #create_periods_summary(doc, periods_db_path, selected_periods)

    # Подключение к базам данных
    periods_conn = sqlite3.connect(periods_db_path)
    logger_stats_conn = sqlite3.connect(logger_stats_db_path)
    
    periods_cursor = periods_conn.cursor()
    logger_stats_cursor = logger_stats_conn.cursor()

    # Инициализация переменных для хранения экстремальных значений
    min_temp = float('inf')
    max_temp = float('-inf')
    min_humidity = float('inf')
    max_humidity = float('-inf')
    min_temp_logger = max_temp_logger = min_humidity_logger = max_humidity_logger = None

    # Номера логгеров с выходом температуры за пределы
    out_of_range_loggers_temp = set()

    # Для расчёта средних значений по зоне и стабильности
    all_temp_avgs = []
    all_hum_avgs = []
    temp_stabilities = []   # (стабильность, номер логгера)
    hum_stabilities = []    # (стабильность, номер логгера)
    # Получение экстремальных значений для всех выбранных периодов
    for period in selected_periods:
        period_id = period[0]
        
        # Получение минимальной и максимальной температуры
        logger_stats_cursor.execute("""
            SELECT MIN(min_value), MAX(max_value), 
                   (SELECT logger_number FROM logger_stats WHERE period_id = ? AND data_type = 'temperature' ORDER BY min_value ASC LIMIT 1),
                   (SELECT logger_number FROM logger_stats WHERE period_id = ? AND data_type = 'temperature' ORDER BY max_value DESC LIMIT 1)
            FROM logger_stats 
            WHERE period_id = ? AND data_type = 'temperature'
        """, (period_id, period_id, period_id))
        temp_data = logger_stats_cursor.fetchone()
        
        if temp_data[0] is not None and temp_data[0] < min_temp:
            min_temp = temp_data[0]
            min_temp_logger = temp_data[2]
        if temp_data[1] is not None and temp_data[1] > max_temp:
            max_temp = temp_data[1]
            max_temp_logger = temp_data[3]

        # Получение минимальной и максимальной влажности
        logger_stats_cursor.execute("""
            SELECT MIN(min_value), MAX(max_value),
                   (SELECT logger_number FROM logger_stats WHERE period_id = ? AND data_type = 'humidity' ORDER BY min_value ASC LIMIT 1),
                   (SELECT logger_number FROM logger_stats WHERE period_id = ? AND data_type = 'humidity' ORDER BY max_value DESC LIMIT 1)
            FROM logger_stats 
            WHERE period_id = ? AND data_type = 'humidity'
        """, (period_id, period_id, period_id))
        humidity_data = logger_stats_cursor.fetchone()
        
        if humidity_data[0] is not None and humidity_data[0] < min_humidity:
            min_humidity = humidity_data[0]
            min_humidity_logger = humidity_data[2]
        if humidity_data[1] is not None and humidity_data[1] > max_humidity:
            max_humidity = humidity_data[1]
            max_humidity_logger = humidity_data[3]

    
    # Добавляем разрыв страницы после таблицы
    doc.add_page_break()

    # Меняем ориентацию на альбомную только для страниц с таблицами
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(297)
    section.page_height = Mm(210)
    section.top_margin = Mm(20)  # 2 см сверху
    section.bottom_margin = Mm(20)  # 2 см снизу (для симметрии)
    section.left_margin = Mm(15)  # 1.5 см слева
    section.right_margin = Mm(15)  # 1.5 см справа

    # Добавляем надпись "7.1 Результаты картирования"
    results_heading = doc.add_paragraph()
    results_heading.add_run().add_tab()  # Добавляем табуляцию
    run = results_heading.add_run("7.1 Результаты картирования")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    results_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Создание таблиц для каждого выбранного периода
    for period in selected_periods:
        period_id = period[0]

        # Получение данных о периоде
        periods_cursor.execute("""
            SELECT start_time, end_time
            FROM periods 
            WHERE id = ?
        """, (period_id,))
        period_data = periods_cursor.fetchone()
        start_time, end_time = period_data

        # Преобразование формата даты и времени
        start_time = datetime.strptime(start_time, "%Y-%m-%d %H:%M:%S").strftime("%d.%m.%Y %H:%M")
        end_time = datetime.strptime(end_time, "%Y-%m-%d %H:%M:%S").strftime("%d.%m.%Y %H:%M")

        # Создание таблиц для температуры и влажности (только если влажность учитывается)
        data_types = ['temperature']
        if humidity_min is not None and humidity_max is not None:
            data_types.append('humidity')
        
        for data_type in data_types:
            # Получение данных логгеров
            logger_stats_cursor.execute("""
                SELECT logger_number, min_value, max_value, avg_value, logger_type
                FROM logger_stats 
                WHERE period_id = ? AND data_type = ?
                ORDER BY 
                    CASE 
                        WHEN logger_type = 'internal' THEN 1
                        ELSE 2
                    END,
                    logger_number
            """, (period_id, data_type))
            logger_data = logger_stats_cursor.fetchall()

            if not logger_data:
                continue  # Пропускаем, если нет данных для этого типа

            # Заголовок таблицы
            # Заголовок таблицы
            unit = "°C" if data_type == 'temperature' else "%"
            header_text = ('Таблица 2 – Результаты картирования по температуре в зоне хранения лекарственных средств' 
                        if data_type == 'temperature' 
                        else 'Таблица 3 – Результаты картирования по относительной влажности в зоне хранения лекарственных средств')

            new_paragraph = doc.add_paragraph()
            run = new_paragraph.add_run(header_text)
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
            new_paragraph.paragraph_format.space_before = Pt(6)  # 6 пт над текстом
            new_paragraph.paragraph_format.space_after = Pt(2)   # 2 пт под текстом
        
            # Определение количества строк и столбцов
            rows = len(logger_data) + 5  # +3 для заголовков и +2 для дополнительной информации
            cols = 8  # Количество столбцов в таблице

            # Добавление таблицы в документ
            table = doc.add_table(rows=rows, cols=cols)
            table.style = 'Table Grid'

            data_type1 = "температура" if data_type == 'temperature' else "относительная влажность"

            # Заголовки таблицы
            headers = [
                "№ логгера", f"Минимальная {data_type1}, {unit}", f"Максимальная {data_type1}, {unit}", 
                f"Средняя {data_type1}, {unit}", "Соответствие критерию", "", "Выполнил", "Дата"
            ]
            for i, header in enumerate(headers):
                cell = table.cell(0, i)
                cell.text = header
                cell.paragraphs[0].runs[0].bold = True  # Делаем текст жирным
                cell.paragraphs[0].runs[0].font.size = Pt(11)  # Устанавливаем размер шрифта
                center_cell_text(cell)

            table.cell(0, 4).merge(table.cell(0, 5))

            subheaders = ["да", "нет"]
            for i, subheader in enumerate(subheaders):
                cell = table.cell(1, i + 4)  # i + 4, потому что "да" и "нет" находятся в 5-м и 6-м столбцах
                cell.text = subheader
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].runs[0].font.size = Pt(11)
                center_cell_text(cell)
            # Объединяем остальные ячейки по вертикали
            for i in [0, 1, 2, 3, 6, 7]:
                table.cell(0, i).merge(table.cell(1, i))

            # Находим минимальное и максимальное значения
            min_value = min(data[1] for data in logger_data)
            max_value = max(data[2] for data in logger_data)

            # Заполнение данными логгеров
            for row, (logger_number, min_val, max_val, avg_value, logger_type) in enumerate(logger_data, start=2):
                table.cell(row, 0).text = str(logger_number)
                if data_type == 'temperature':
                    if logger_type == 'internal':
                        if min_val < temp_min or max_val > temp_max:
                            out_of_range_loggers_temp.add(str(logger_number))

                # Копим данные для расчёта средних по зоне и стабильности
                if logger_type == 'internal':
                    if data_type == 'temperature':
                        all_temp_avgs.append(avg_value)
                        stability = max(max_val - avg_value, avg_value - min_val)
                        temp_stabilities.append((stability, logger_number))
                    elif data_type == 'humidity':
                        all_hum_avgs.append(avg_value)
                        stability = max(max_val - avg_value, avg_value - min_val)
                        hum_stabilities.append((stability, logger_number))
                for col in range(8):
                    cell = table.cell(row, col)
                    if col == 1:
                        cell.text = f"{min_val:.2f}"
                        if min_val == min_value:
                            set_cell_background(cell, "87CEFA")
                    elif col == 2:
                        cell.text = f"{max_val:.2f}"
                        if max_val == max_value:
                            set_cell_background(cell, "FFB6C1")
                    elif col == 3:
                        cell.text = f"{avg_value:.2f}"
                    elif logger_type == 'internal':
                        if col == 4:
                            if data_type == 'temperature':
                                if temp_min <= min_val <= max_val <= temp_max:
                                    cell.text = "+"
                            elif data_type == 'humidity' and humidity_min is not None and humidity_max is not None:
                                if humidity_min <= min_val <= max_val <= humidity_max:
                                    cell.text = "+"
                        elif col == 5:
                            if data_type == 'temperature':
                                if min_val < temp_min or max_val > temp_max:
                                    cell.text = "+"
                            elif data_type == 'humidity' and humidity_min is not None and humidity_max is not None:
                                if min_val < humidity_min or max_val > humidity_max:
                                    cell.text = "+"
                    
                    
                    # Центрирование текста в ячейке
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Объединение ячеек "Выполнил" и "Дата"
            table.cell(2, 6).text = "Сенчина Т.А."
            center_cell_text(table.cell(2, 6))
            for row in range(3, len(logger_data) + 2):
                table.cell(2, 6).merge(table.cell(row, 6))
            
            table.cell(2, 7).text = contract_date
            center_cell_text(table.cell(2, 7))
            for row in range(3, len(logger_data) + 2):
                table.cell(2, 7).merge(table.cell(row, 7))

            # Добавление нижних строк
            start_row = len(logger_data) + 2
            table.cell(start_row, 0).text = "Начало периода картирования (дата, время)"
            table.cell(start_row, 0).merge(table.cell(start_row, 3))
            table.cell(start_row, 4).text = start_time
            table.cell(start_row, 4).merge(table.cell(start_row, 7))

            table.cell(start_row + 1, 0).text = "Окончание периода картирования (дата, время)"
            table.cell(start_row + 1, 0).merge(table.cell(start_row + 1, 3))
            table.cell(start_row + 1, 4).text = end_time
            table.cell(start_row + 1, 4).merge(table.cell(start_row + 1, 7))

            table.cell(start_row + 2, 0).text = "Проверено"
            table.cell(start_row + 2, 0).merge(table.cell(start_row + 2, 0))
            table.cell(start_row + 2, 1).text = "Харьевич А.И."
            table.cell(start_row + 2, 1).merge(table.cell(start_row + 2, 3))
            table.cell(start_row + 2, 4).text = "Дата"
            table.cell(start_row + 2, 4).merge(table.cell(start_row + 2, 5))
            table.cell(start_row + 2, 6).text = contract_date
            table.cell(start_row + 2, 6).merge(table.cell(start_row + 2, 7))

            

           # Установка общей ширины таблицы и относительной ширины столбцов
            table.autofit = False
            table.width = Inches(7.5)  # Примерная ширина страницы A4 в альбомной ориентации
            column_widths = [10, 15, 15, 15, 10, 10, 15, 10]  # Относительные ширины в процентах

            for i, width in enumerate(column_widths):
                table.columns[i].width = Inches(7.5 * width / 100)

            # Центрирование всех ячеек по вертикали
            for row in table.rows:
                for cell in row.cells:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # Обновляем индекс параграфа
            paragraph_index = len(doc.paragraphs) - 1


    # Закрытие соединений с базами данных
    periods_conn.close()
    logger_stats_conn.close()

    # Добавляем разрыв страницы после таблицы
    doc.add_page_break()

    # Возвращаем ориентацию на портретную для следующих страниц
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(15)  # 2 см сверху
    section.bottom_margin = Mm(15)  # 2 см снизу (для симметрии)
    section.left_margin = Mm(20)  # 1.5 см слева
    section.right_margin = Mm(20)  # 1.5 см справа

    # Добавляем надпись "8. Анализ результатов картирования"
    results_heading = doc.add_paragraph()
    results_heading.add_run().add_tab()  # Добавляем табуляцию
    run = results_heading.add_run("8. Анализ результатов картирования")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    results_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    results_heading.paragraph_format.space_after = Pt(6)

    # Добавление текста перед таблицей
    plan_text = doc.add_paragraph()
    plan_text.add_run().add_tab()
    run = plan_text.add_run("Результаты проведения исследований представлены в таблице 4.")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    # Добавление заголовка таблицы
    table_heading = doc.add_paragraph()
    run = table_heading.add_run("Таблица 4 – Результаты проведения исследований")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    table_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    table_heading.paragraph_format.space_before = Pt(4)  # 4 пт над текстом
    table_heading.paragraph_format.space_after = Pt(2)   # 2 пт под текстом

    # Определение отклонений от нормы (для холодной/горячей точки)
    cold_deviation = abs(min_temp - temp_min) if temp_min is not None else float('inf')
    hot_deviation = abs(max_temp - temp_max) if temp_max is not None else float('inf')

    # Добавьте эту функцию, если её еще нет
    def set_cell_alignment(cell, horizontal_alignment, vertical_alignment):
        cell.vertical_alignment = vertical_alignment
        for paragraph in cell.paragraphs:
            paragraph.alignment = horizontal_alignment
    
    # Формирование строки по логгерам с превышением температуры
    out_of_range_result_temp = (
        f"№ логгера – {', '.join(sorted(out_of_range_loggers_temp))}"
        if out_of_range_loggers_temp
        else "Не выявлены"
    )

    # Средняя температура и влажность по всей зоне
    avg_temp_zone = sum(all_temp_avgs) / len(all_temp_avgs) if all_temp_avgs else None
    avg_hum_zone = sum(all_hum_avgs) / len(all_hum_avgs) if all_hum_avgs else None

    # Однородность как размах (max - min) по внутренним логгерам
    homogeneity_temp = (
        max_temp - min_temp
        if min_temp != float('inf') and max_temp != float('-inf')
        else None
    )
    homogeneity_hum = (
        max_humidity - min_humidity
        if min_humidity != float('inf') and max_humidity != float('-inf')
        else None
    )

    # Стабильность как максимальное отклонение от среднего по логгеру
    stab_temp_value = stab_temp_logger = None
    if temp_stabilities:
        stab_temp_value, stab_temp_logger = max(temp_stabilities, key=lambda x: x[0])

    stab_hum_value = stab_hum_logger = None
    if hum_stabilities:
        stab_hum_value, stab_hum_logger = max(hum_stabilities, key=lambda x: x[0])

    # Подготовка текстовых значений
    avg_temp_text = f"{avg_temp_zone:.2f}" if avg_temp_zone is not None else "—"
    avg_hum_text = f"{avg_hum_zone:.2f}" if avg_hum_zone is not None else "—"

    hom_temp_text = f"{homogeneity_temp:.2f}" if homogeneity_temp is not None else "—"
    hom_hum_text = f"{homogeneity_hum:.2f}" if homogeneity_hum is not None else "—"

    stab_temp_text = (
        f"{stab_temp_value:.2f} (№ логгера – {stab_temp_logger})"
        if stab_temp_logger is not None
        else "—"
    )
    stab_hum_text = (
        f"{stab_hum_value:.2f} (№ логгера – {stab_hum_logger})"
        if stab_hum_logger is not None
        else "—"
    )

    min_temp_text = (
        f"{min_temp:.2f} (№ логгера – {min_temp_logger})"
        if min_temp != float('inf') and min_temp_logger is not None
        else "—"
    )
    max_temp_text = (
        f"{max_temp:.2f} (№ логгера – {max_temp_logger})"
        if max_temp != float('-inf') and max_temp_logger is not None
        else "—"
    )
    min_hum_text = (
        f"{min_humidity:.2f} (№ логгера – {min_humidity_logger})"
        if min_humidity != float('inf') and min_humidity_logger is not None
        else "—"
    )
    max_hum_text = (
        f"{max_humidity:.2f} (№ логгера – {max_humidity_logger})"
        if max_humidity != float('-inf') and max_humidity_logger is not None
        else "—"
    )

    # Холодная и горячая точки – просто логгеры с минимальной и максимальной температурой
    cold_point_text = f"{min_temp_logger}" if min_temp_logger is not None else "Не выявлено"
    hot_point_text = f"{max_temp_logger}" if max_temp_logger is not None else "Не выявлено"

    # Если из other_info передан текст однородности, используем его вместо рассчитанного по БД
    if temp_homogeneity_text is not None:
        hom_temp_text = temp_homogeneity_text
    if hum_homogeneity_text is not None:
        hom_hum_text = hum_homogeneity_text

    # Определение данных для новой таблицы 4
    data = [
        ("Наименование показателя", "Результат"),
        (
            "Результаты картирования",
            f"{mapping_results}\nСм. табл. 2, 3, Приложение 4, 6."
        ),
        (
            "Области с температурой, выходящей за допустимые граничные значения, № логгера(-ов)",
            out_of_range_result_temp
        ),
        ("Минимальная температура, °C", min_temp_text),
        ("Максимальная температура, °C", max_temp_text),
        ("Средняя температура (для всей зоны хранения), °C", avg_temp_text),
        ("Расположение холодной точки, № логгера", cold_point_text),
        ("Расположение горячей точки, № логгера", hot_point_text),
        ("Однородность температуры (максимальное значение), °С", hom_temp_text),
        ("Стабильность температуры (максимальное значение), °С", stab_temp_text),
    ]
    
    # Добавляем строки влажности только если влажность учитывается
    if humidity_min is not None and humidity_max is not None:
        data.extend([
            ("Минимальная относительная влажность, %", min_hum_text),
            ("Максимальная относительная влажность, %", max_hum_text),
            ("Средняя относительная влажность (для всей зоны хранения), %", avg_hum_text),
            ("Однородность влажности (максимальное значение), %", hom_hum_text),
            ("Стабильность влажности (максимальное значение), %", stab_hum_text),
        ])

    # Создание таблицы
    if data:
        table = doc.add_table(rows=len(data), cols=2)
        table.style = 'Table Grid'

        # Заполнение таблицы
        for i, (key, value) in enumerate(data):
            cell_left = table.cell(i, 0)
            cell_right = table.cell(i, 1)
            
            cell_left.text = key
            cell_right.text = value
            
            if i == 0:
                # Центрирование текста в первой строке
                set_cell_alignment(cell_left, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_VERTICAL.CENTER)
                set_cell_alignment(cell_right, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_VERTICAL.CENTER)
                
                # Жирный шрифт для первой строки
                cell_left.paragraphs[0].runs[0].bold = True
                cell_right.paragraphs[0].runs[0].bold = True
            else:
                # Выравнивание по левому краю для остальных строк
                set_cell_alignment(cell_left, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_VERTICAL.CENTER)
                set_cell_alignment(cell_right, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_VERTICAL.CENTER)
            
            # Жирный шрифт для всех ячеек в первом столбце
            cell_left.paragraphs[0].runs[0].bold = True

        # Установка ширины столбцов
        table.columns[0].width = Inches(3)
        table.columns[1].width = Inches(4.5)
    else:
        print("Внимание: Нет данных для создания таблицы.")

     # Добавление заключения
    conclusion_heading = doc.add_paragraph()
    conclusion_heading.paragraph_format.first_line_indent = Inches(0.5)  # Добавляем табуляцию
    run = conclusion_heading.add_run("Заключение:")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    conclusion_text = doc.add_paragraph()
    conclusion_text.add_run().add_tab()
    run = conclusion_text.add_run(conclusion)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    # Рекомендации (сразу после заключения: 1) 2), интервал 0, отступ как у заголовка)
    recs = selected_recommendations or []
    if recs:
        rec_heading = doc.add_paragraph()
        rec_heading.paragraph_format.first_line_indent = Inches(0.5)
        run_h = rec_heading.add_run("Рекомендации:")
        run_h.bold = True
        run_h.font.name = 'Times New Roman'
        run_h.font.size = Pt(12)
        rec_heading.paragraph_format.space_before = Pt(6)
        rec_heading.paragraph_format.space_after = Pt(0)
        for i, rec in enumerate(recs, 1):
            p = doc.add_paragraph()
            r = p.add_run(f"{i}) {rec}")
            r.font.name = 'Times New Roman'
            r.font.size = Pt(12)
            p.paragraph_format.first_line_indent = Inches(0.5)
            p.paragraph_format.space_after = Pt(0)

    # Добавляем разрыв страницы после таблицы
    doc.add_page_break()

    return paragraph_index, out_of_range_result_temp

#def create_periods_summary(doc, periods_db_path, selected_periods):
    # Подключение к базе данных периодов
    periods_conn = sqlite3.connect(periods_db_path)
    periods_cursor = periods_conn.cursor()

    # Добавление заголовка "6 Картирование"
    heading_6 = doc.add_paragraph()
    heading_6.add_run().add_tab()
    run = heading_6.add_run("6 Картирование")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    heading_6.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Добавление текста перед таблицей
    plan_text = doc.add_paragraph()
    plan_text.add_run().add_tab()
    run = plan_text.add_run("План-график картирования представлен в таблице 1.")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    # Добавление заголовка таблицы
    table_heading = doc.add_paragraph()
    run = table_heading.add_run("Таблица 1 – План-график картирования")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    table_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Создание таблицы
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    # Установка ширины столбцов (в процентах от ширины страницы)
    widths = [20, 30, 30, 20]
    for i, width in enumerate(widths):
        table.columns[i].width = Cm(width * 0.15)

    # Заголовки таблицы
    headers = ["Дата / время проведения", "Исследование", "Диапазон климатических условий (температура, влажность)", "Время проведения исследований"]
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(header)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Заполнение таблицы данными выбранных периодов
    for index, period in enumerate(selected_periods, start=1):
        period_id = period[0]
        periods_cursor.execute("""
            SELECT start_time, end_time, required_mode_from, required_mode_to
            FROM periods 
            WHERE id = ?
        """, (period_id,))
        period_data = periods_cursor.fetchone()
        
        row = table.add_row().cells
        
        # Форматирование даты и времени
        start = datetime.strptime(period_data[0], "%Y-%m-%d %H:%M:%S")
        end = datetime.strptime(period_data[1], "%Y-%m-%d %H:%M:%S")
        row[0].text = f"{start.strftime('%d.%m.%Y %H:%M')} - {end.strftime('%d.%m.%Y %H:%M')}"
        
        row[1].text = "Здесь будет ввод пользователя"  # Пока нет данных об исследовании
        
        # Форматирование диапазона климатических условий
        row[2].text = f"Температура: +{period_data[2]}℃...+{period_data[3]}℃"
        
        # Расчет продолжительности периода
        duration = end - start
        days = duration.days
        hours, remainder = divmod(duration.seconds, 3600)
        minutes, _ = divmod(remainder, 60)
        
        duration_str = ""
        if days > 0:
            duration_str += f"{days} {'день' if days == 1 else 'дня' if 1 < days < 5 else 'дней'}"
        if hours > 0:
            duration_str += f" {hours} {'час' if hours == 1 else 'часа' if 1 < hours < 5 else 'часов'}"
        if minutes > 0:
            duration_str += f" {minutes} {'минута' if minutes == 1 else 'минуты' if 1 < minutes < 5 else 'минут'}"
        
        row[3].text = duration_str.strip()

        # Центрирование текста в ячейках
        for cell in row:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Закрытие соединения с базой данных
    periods_conn.close()

    # Добавляем разрыв страницы после таблицы
    doc.add_page_break()


