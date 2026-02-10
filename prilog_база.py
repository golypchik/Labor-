import docx 
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt, Cm, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
from copy import deepcopy
import sys 
from docx.shared import Mm
import sys
import os

def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    try:
        # PyInstaller creates a temp folder and stores path in _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")

    path = os.path.join(base_path, relative_path)
    
    if not os.path.exists(path):
        # If the file is not found in the PyInstaller temp folder, try the current directory
        path = os.path.join(os.path.dirname(os.path.abspath(__file__)), relative_path)
    
    if not os.path.exists(path):
        # If still not found, try the directory of the executable
        path = os.path.join(os.path.dirname(sys.executable), relative_path)
    
    return path

def create_appendices(doc,images,saved_risk_areas,selected_template, period_images):
    
    # Приложение 1: Планировка зоны хранения
    p = doc.add_paragraph()
    run = p.add_run("Приложение 1")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


    
    p = doc.add_paragraph()
    run = p.add_run("Планировка зоны хранения лекарственных средств")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Добавление изображения планировки
    if images['layout']:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)  # Добавляем 6 пт перед изображением
        p.paragraph_format.space_after = Pt(6)   # Добавляем 6 пт после изображения
        run = p.add_run()
        run.add_picture(images['layout'], width=Inches(6))  # Увеличенный размер
    else:
        doc.add_paragraph("Изображение планировки не загружено")

    # Добавление разрыва страницы
    doc.add_page_break()

    # Приложение 2: Анализ рисков
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(297)
    section.page_height = Mm(210)
    section.top_margin = Mm(20)  # 2 см сверху
    section.bottom_margin = Mm(20)  # 2 см снизу (для симметрии)
    section.left_margin = Mm(15)  # 1.5 см слева
    section.right_margin = Mm(15)  # 1.5 см справа


    p = doc.add_paragraph()
    run = p.add_run("Приложение 2")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    run = p.add_run("Анализ рисков")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Добавление пустой строки
    doc.add_paragraph()


    
    table = doc.add_table(rows=2, cols=9)
    table.style = 'Table Grid'
    table.autofit = False
    table.allow_autofit = False

    # Установка ширины столбцов
    widths = [0.3, 2, 1, 1, 1, 2, 1, 1, 1]
    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width)

    # Заголовки таблицы
    headers = ["№", "Идентифицированный риск", "Последствия возникнове-ния риска (P)", "Вероятность возникнове-ния риска (S)", 
            "Оценка риска (OR)", "Меры снижения", "Остаточный риск"]
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(9.5)
        cell.width = Inches(widths[i])

    # Объединение ячеек в первых 6 столбцах
    for i in range(6):
        table.cell(0, i).merge(table.cell(1, i))

    # Объединение ячеек для "Остаточный риск" по горизонтали
    table.cell(0, 6).merge(table.cell(0, 8))

    # Добавление подзаголовков для "Остаточный риск"
    table.cell(1, 6).text = "Последствия возникнове-ния риска (P)"
    table.cell(1, 7).text = "Вероятность возникнове-ния риска (S)"
    table.cell(1, 8).text = "Оценка риска (OR)"
    for i in range(6, 9):
        cell = table.cell(1, i)
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(9.5)

    # Добавление данных в таблицу
    for i, risk in enumerate(saved_risk_areas, start=1):
        row = table.add_row()
        row.cells[0].text = str(i)
        if selected_template.get() == "ХОЛОДИЛЬНИК(БЕЗ ОТКРЫТИЯ)":
            row.cells[1].text = f"Риск выхода климатических условий (температуры) за установленные пределы в зоне хранения лекарственных средств {risk}"
        elif  selected_template.get() in ["ОБЪЕКТ ХРАНЕНИЯ ЛЕКАРСТВЕННЫХ СРЕДСТВ", "ЗОНА ХРАНЕНИЯ ЛЕКАРСТВЕННЫХ СРЕДСТВ"]:
            row.cells[1].text = f"Риск выхода климатических условий (температуры, влажности) за установленные пределы в зоне хранения лекарственных средств {risk}"
        row.cells[2].text = "3"
        row.cells[3].text = "3"
        row.cells[4].text = "9"
        row.cells[5].text = f"Провести анализ зоны хранения лекарственных средств {risk} и на основании этого анализа обозначить места установки логгеров на схеме размещения"
        row.cells[6].text = "2"
        row.cells[7].text = "1"
        row.cells[8].text = "2"

        # Центрирование текста и выделение жирным шрифтом в ячейках
        for j, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(9.5)
            if j not in [1, 5]:  # Исключаем ячейки с индексами 1 и 5
                for run in paragraph.runs:
                    run.bold = True

    # После заполнения таблицы данными, еще раз устанавливаем ширину
    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width)

    # Принудительно применяем ширину к каждой ячейке
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(widths[i])

    # Добавление разрыва страницы
    doc.add_page_break()

    # Вставка содержимого из rrr.docx сразу после таблицы рисков
    rrr_path = resource_path('rrr.docx')
    if os.path.exists(rrr_path):
        rrr_doc = Document(rrr_path)
        
        # Находим индекс последнего элемента таблицы рисков
        insert_index = -1
        for i, element in enumerate(doc.element.body):
            if isinstance(element, docx.oxml.table.CT_Tbl):
                insert_index = i + 1
        
        # Добавляем разрыв страницы после последней таблицы
        if insert_index != -1:
            page_break = OxmlElement('w:p')
            run = OxmlElement('w:r')
            br = OxmlElement('w:br')
            br.set(qn('w:type'), 'page')
            run.append(br)
            page_break.append(run)
            doc.element.body.insert(insert_index, page_break)
            insert_index += 1
        
        # Копируем содержимое из rrr.docx
        for i, element in enumerate(rrr_doc.element.body):
            new_element = deepcopy(element)
            doc.element.body.insert(insert_index + i, new_element)
        
        # Изменяем интервал для всех параграфов в новой секции
        for paragraph in doc.paragraphs[insert_index:insert_index+len(rrr_doc.paragraphs)]:
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
    else:
        print(f"Файл {rrr_path} не найден")

    # Возвращаем ориентацию на портретную для следующих страниц
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(15)  # 2 см сверху
    section.bottom_margin = Mm(15)  # 2 см снизу (для симметрии)
    section.left_margin = Mm(20)  # 1.5 см слева
    section.right_margin = Mm(20)  # 1.5 см справа

    

    p = doc.add_paragraph()
    run = p.add_run("Приложение 3")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    run = p.add_run("Схема размещения логгеров")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


    # Добавление изображения планировки
    if images['loggers']:
        # Создаем параграф для изображения и центрируем его
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)  # Добавляем 6 пт перед изображением
        p.paragraph_format.space_after = Pt(6)   # Добавляем 6 пт после изображения
        run = p.add_run()
        # Устанавливаем ширину изображения в 6 дюймов (около 15 см)
        run.add_picture(images['loggers'], width=Inches(6))
    else:
        doc.add_paragraph("Изображение планировки не загружено")

    # Добавление разрыва страницы
    doc.add_page_break()

    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(297)
    section.page_height = Mm(210)
    section.top_margin = Mm(20)  # 2 см сверху
    section.bottom_margin = Mm(20)  # 2 см снизу (для симметрии)
    section.left_margin = Mm(15)  # 1.5 см слева
    section.right_margin = Mm(15)  # 1.5 см справа

    p = doc.add_paragraph()
    run = p.add_run("Приложение 4")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if selected_template.get() == "ХОЛОДИЛЬНИК(БЕЗ ОТКРЫТИЯ)":
        p = doc.add_paragraph()
        run = p.add_run("Графики распределения температуры  при проведении исследовании")
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Добавление пустой строки
        doc.add_paragraph()


        # Добавление заголовка таблицы
        table_heading = doc.add_paragraph()
        run = table_heading.add_run("На рисунке 4.1 представлен график распределения температуры в зоне хранения лекарственных средств (в холодильнике) при режиме тер-морегуляции №1 на протяжении всего времени исследования")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        table_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        
        # График температуры для первого периода
        if period_images and 1 in period_images and 'temp_fridge' in period_images[1]:
            # Создаем параграф для изображения и центрируем его
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)  # Добавляем 6 пт перед изображением
            p.paragraph_format.space_after = Pt(6)   # Добавляем 6 пт после изображения
            run = p.add_run()
            # Устанавливаем ширину изображения в 6 дюймов (около 15 см)
            run.add_picture(period_images[1]['temp_fridge'], width=Inches(9))
        else:
            doc.add_paragraph("График температуры для первого периода не загружен")


        # Добавление подписи к рисунку
        caption = doc.add_paragraph("Рисунок 4.1 – График распределения температуры в зоне хранения лекарственных средств (в холодильнике) при режиме терморегуляции №1 на протяжении всего времени исследования")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in caption.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
    
    elif selected_template.get() == "ОБЪЕКТ ХРАНЕНИЯ ЛЕКАРСТВЕННЫХ СРЕДСТВ" :
        p = doc.add_paragraph()
        run = p.add_run("Графики распределения температуры и влажности при проведении исследований")
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Добавление пустой строки
        doc.add_paragraph()

        # Добавление заголовка таблицы
        table_heading = doc.add_paragraph()
        run = table_heading.add_run("На рисунке 4.1 представлен график распределения температуры в зоне хранения лекарственных средств на протяжении всего времени исследования")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        table_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT


        # График температуры для первого периода
        if period_images and 1 in period_images and 'temp_loggers' in period_images[1]:
            # Создаем параграф для изображения и центрируем его
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)  # Добавляем 6 пт перед изображением
            p.paragraph_format.space_after = Pt(6)   # Добавляем 6 пт после изображения
            run = p.add_run()
            # Устанавливаем ширину изображения в 6 дюймов (около 15 см)
            run.add_picture(period_images[1]['temp_loggers'], width=Inches(9))
        else:
            doc.add_paragraph("График температуры для первого периода не загружен")

   
        # Добавление подписи к рисунку
        caption = doc.add_paragraph("Рисунок 4.1 – График распределения температуры в зоне хранения лекарственных средств на протяжении всего времени исследования")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in caption.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

        # Добавление разрыва страницы
        doc.add_page_break()

        

        # Добавление заголовка таблицы
        table_heading = doc.add_paragraph()
        run = table_heading.add_run("На рисунке 4.2 представлен график распределения относительной влажности в зоне хранения лекарственных средств на протяжении всего времени исследования")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        table_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT


        # График температуры для первого периода
        if period_images and 1 in  period_images and 'humidity_loggers' in  period_images[1]:
            # Создаем параграф для изображения и центрируем его
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)  # Добавляем 6 пт перед изображением
            p.paragraph_format.space_after = Pt(6)   # Добавляем 6 пт после изображения
            run = p.add_run()
            # Устанавливаем ширину изображения в 6 дюймов (около 15 см)
            run.add_picture( period_images[1]['humidity_loggers'], width=Inches(9))
        else:
            doc.add_paragraph("График температуры для первого периода не загружен")


        # Добавление подписи к рисунку
        caption = doc.add_paragraph("Рисунок 4.2 – График распределения относительной влажности в зоне хранения лекарственных средств на протяжении всего времени исследования ")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in caption.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

        

    elif selected_template.get() == "ЗОНА ХРАНЕНИЯ ЛЕКАРСТВЕННЫХ СРЕДСТВ" :
        p = doc.add_paragraph()
        run = p.add_run("Графики распределения температуры и влажности при проведении исследований")
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Добавление пустой строки
        doc.add_paragraph()

        # Добавление заголовка таблицы
        table_heading = doc.add_paragraph()
        run = table_heading.add_run("На рисунке 4.1 представлен график распределения температуры в зоне хранения лекарственных средств на протяжении всего времени исследования")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        table_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT


        # График температуры для первого периода
        if period_images and 1 in period_images and 'temp_loggers' in period_images[1]:
            # Создаем параграф для изображения и центрируем его
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)  # Добавляем 6 пт перед изображением
            p.paragraph_format.space_after = Pt(6)   # Добавляем 6 пт после изображения
            run = p.add_run()
            # Устанавливаем ширину изображения в 6 дюймов (около 15 см)
            run.add_picture(period_images[1]['temp_loggers'], width=Inches(9))
        else:
            doc.add_paragraph("График температуры для первого периода не загружен")



        # Добавление подписи к рисунку
        caption = doc.add_paragraph("Рисунок 4.1 – График распределения температуры в зоне хранения лекарственных средств на протяжении всего времени исследования")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in caption.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

        # Добавление разрыва страницы
        doc.add_page_break()

        # Добавление заголовка таблицы
        table_heading = doc.add_paragraph()
        run = table_heading.add_run("На рисунке 4.2 представлен график распределения температуры около хоны хранения лекарственных средств на протяжении всего времени исследования")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        table_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT


        # График температуры для первого периода
        if period_images and 1 in  period_images and 'temp_external' in  period_images[1]:
            # Создаем параграф для изображения и центрируем его
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)  # Добавляем 6 пт перед изображением
            p.paragraph_format.space_after = Pt(6)   # Добавляем 6 пт после изображения
            run = p.add_run()
            # Устанавливаем ширину изображения в 6 дюймов (около 15 см)
            run.add_picture( period_images[1]['temp_external'], width=Inches(9))
        else:
            doc.add_paragraph("График температуры для первого периода не загружен")



        # Добавление подписи к рисунку
        caption = doc.add_paragraph("Рисунок 4.2 – График распределения температуры около зоны хранения лекарственных средств на протяжении всего времени исследования")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in caption.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

        # Добавление разрыва страницы
        doc.add_page_break()

        # Добавление заголовка таблицы
        table_heading = doc.add_paragraph()
        run = table_heading.add_run("На рисунке 4.3 представлен график распределения относительной влажности в зоне хранения лекарственных средств на протяжении всего времени исследования")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        table_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT


        # График температуры для первого периода
        if period_images and 1 in  period_images and 'humidity_loggers' in  period_images[1]:
            # Создаем параграф для изображения и центрируем его
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)  # Добавляем 6 пт перед изображением
            p.paragraph_format.space_after = Pt(6)   # Добавляем 6 пт после изображения
            run = p.add_run()
            # Устанавливаем ширину изображения в 6 дюймов (около 15 см)
            run.add_picture( period_images[1]['humidity_loggers'], width=Inches(9))
        else:
            doc.add_paragraph("График температуры для первого периода не загружен")


        # Добавление подписи к рисунку
        caption = doc.add_paragraph("Рисунок 4.3 – График распределения относительной влажности в зоне хранения лекарственных средств на протяжении всего времени исследования ")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in caption.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

        # Добавление разрыва страницы
        doc.add_page_break()

        # Добавление заголовка таблицы
        table_heading = doc.add_paragraph()
        run = table_heading.add_run("На рисунке 4.4 представлен график распределения относительной влажности около зоны хранения лекарственных средств на протяжении всего времени исследования")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        table_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT


        # График температуры для первого периода
        if period_images and 1 in  period_images and 'humidity_external' in  period_images[1]:
            # Создаем параграф для изображения и центрируем humidity_external
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run()
            # Устанавливаем ширину изображения в 6 дюймов (около 15 см)
            run.add_picture( period_images[1]['humidity_external'], width=Inches(9))
        else:
            doc.add_paragraph("График температуры для первого периода не загружен")


        # Добавление подписи к рисунку
        caption = doc.add_paragraph("Рисунок 4.4 – График распределения относительной влажности около зоны хранения лекарственных средств на протяжении всего времени исследования ")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in caption.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

    # Добавление разрыва страницы
    doc.add_page_break()

    # Возвращаем ориентацию на портретную для следующих страниц
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(15)  # 2 см сверху
    section.bottom_margin = Mm(15)  # 2 см снизу (для симметрии)
    section.left_margin = Mm(20)  # 1.5 см слева
    section.right_margin = Mm(20)  # 1.5 см справа

    p = doc.add_paragraph()
    run = p.add_run("Приложение 5")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if  selected_template.get() == "ХОЛОДИЛЬНИК(БЕЗ ОТКРЫТИЯ)":
        p = doc.add_paragraph()
        run = p.add_run("Температурная карта")
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Добавление пустой строки
        doc.add_paragraph()

        # Добавление заголовка таблицы
        table_heading = doc.add_paragraph()
        run = table_heading.add_run("На рисунке 5.1 представлена температурная карта в зоне хранения лекарственных средств на протяжении всего времени исследования")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        table_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Добавление изображения планировки
        if  images['temp_map']:
            # Создаем параграф для изображения и центрируем его
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run()
            # Устанавливаем ширину изображения в 6 дюймов (около 15 см)
            run.add_picture( images['temp_map'], width=Inches(9))
        else:
            doc.add_paragraph("Изображение планировки не загружено")


        # Добавление подписи к рисунку
        caption = doc.add_paragraph("Рисунок 5.1 – Температурная карта в зоне хранения лекарственных средств на протяжении все-го времени исследования")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in caption.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

        # Добавление разрыва страницы
        doc.add_page_break()

        
        
    elif  selected_template.get() in ["ОБЪЕКТ ХРАНЕНИЯ ЛЕКАРСТВЕННЫХ СРЕДСТВ", "ЗОНА ХРАНЕНИЯ ЛЕКАРСТВЕННЫХ СРЕДСТВ"]:
        p = doc.add_paragraph()
        run = p.add_run("Температурная и влажностная карты")
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Добавление пустой строки
        doc.add_paragraph()

        # Добавление заголовка таблицы
        table_heading = doc.add_paragraph()
        run = table_heading.add_run("На рисунке 5.1 представлена температурная карта в зоне хранения лекарственных средств на протяжении всего времени исследования")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        table_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT


        # Добавление изображения планировки
        if  images['temp_map']:
            # Создаем параграф для изображения и центрируем его
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run()
            # Устанавливаем ширину изображения в 6 дюймов (около 15 см)
            run.add_picture( images['temp_map'], width=Inches(9))
        else:
            doc.add_paragraph("Изображение планировки не загружено")


        # Добавление подписи к рисунку
        caption = doc.add_paragraph("Рисунок 5.1 – Температурная карта в зоне хранения лекарственных средств на протяжении всего времени исследования")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in caption.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

        # Добавление разрыва страницы
        doc.add_page_break()

        # Добавление заголовка таблицы
        table_heading = doc.add_paragraph()
        run = table_heading.add_run("На рисунке 5.2 представлена влажностная карта в зоне хранения лекарственных средств на протяжении всего времени исследования")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        table_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

 
        # Добавление изображения планировки
        if  images['humidity_map']:
            # Создаем параграф для изображения и центрируем его
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run()
            # Устанавливаем ширину изображения в 6 дюймов (около 15 см)
            run.add_picture( images['humidity_map'], width=Inches(9))
        else:
            doc.add_paragraph("Изображение планировки не загружено")


        # Добавление подписи к рисунку
        caption = doc.add_paragraph("Рисунок 5.2 – Влажностная карта в зоне хранения лекарственных средств на протяжении всего времени исследования ")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in caption.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

