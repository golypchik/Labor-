import docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt, Cm, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
from copy import deepcopy
import sys
from docx.shared import Mm
import sys
import os
from docx.enum.style import WD_STYLE_TYPE
from PIL import Image

def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    try:
        # PyInstaller creates a temp folder and stores path in _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")

    path = os.path.join(base_path, relative_path)

    if not os.path.exists(path):
        # If the file is not found in the PyInstaller temp folder, try the current directory
        path = os.path.join(os.path.dirname(os.path.abspath(__file__)), relative_path)

    return path

def get_image_size_for_page(image_path, is_landscape=False, max_width_inches=None):
    """
    Получает размер изображения с учетом ориентации страницы.
    Если изображение больше допустимого размера, уменьшает его.
    
    Args:
        image_path: путь к изображению
        is_landscape: True если альбомная ориентация, False если книжная
        max_width_inches: максимальная ширина в дюймах (если None, вычисляется автоматически)
    
    Returns:
        tuple: (width_inches, height_inches) - размеры для вставки
    """
    try:
        img = Image.open(image_path)
        original_width_px, original_height_px = img.size
        
        # Определяем максимальные размеры в зависимости от ориентации
        if max_width_inches is None:
            if is_landscape:
                # Альбомная ориентация: ограничения 12 см по высоте и 26 см по ширине
                max_width_inches = 26.0 / 2.54  # 26 см = ~10.24 дюймов
                max_height_inches = 12.0 / 2.54  # 12 см = ~4.72 дюймов
            else:
                # Книжная ориентация: ограничения 23 см по высоте и 16 см по ширине
                max_width_inches = 16.0 / 2.54  # 16 см = ~6.30 дюймов
                max_height_inches = 23.0 / 2.54  # 23 см = ~9.06 дюймов
        else:
            # Если указана максимальная ширина, вычисляем высоту пропорционально
            max_height_inches = max_width_inches * (original_height_px / original_width_px) if original_width_px > 0 else max_width_inches
        
        # Вычисляем размеры с сохранением пропорций
        # Предполагаем DPI = 96 для конвертации пикселей в дюймы
        original_width_inches = original_width_px / 96.0
        original_height_inches = original_height_px / 96.0
        
        # Проверяем, нужно ли уменьшать изображение
        # Если изображение меньше максимальных размеров, оставляем как есть
        if original_width_inches <= max_width_inches and original_height_inches <= max_height_inches:
            return (original_width_inches, original_height_inches)
        
        # Вычисляем коэффициенты масштабирования
        width_ratio = max_width_inches / original_width_inches if original_width_inches > 0 else 1.0
        height_ratio = max_height_inches / original_height_inches if original_height_inches > 0 else 1.0
        
        # Используем меньший коэффициент, чтобы изображение влезло в оба ограничения
        ratio = min(width_ratio, height_ratio, 1.0)  # Не увеличиваем, только уменьшаем
        
        final_width = original_width_inches * ratio
        final_height = original_height_inches * ratio
        
        return (final_width, final_height)
    except Exception as e:
        print(f"Ошибка обработки изображения {image_path}: {e}")
        # Возвращаем размер по умолчанию
        if is_landscape:
            return (9.0, 6.0)
        else:
            return (6.0, 9.0)

def get_fixed_image_size_for_logger_screenshots():
    """
    Возвращает фиксированный размер для скриншотов логгеров в приложении 5.
    Размер: высота 12.10 см, ширина 26 см.
    
    Returns:
        tuple: (width_inches, height_inches) - фиксированные размеры для вставки
    """
    # Конвертируем сантиметры в дюймы (1 дюйм = 2.54 см)
    width_inches = 26.0 / 2.54
    height_inches = 12.10 / 2.54
    return (width_inches, height_inches)

def create_appendices(doc, images, saved_risk_areas, selected_template, period_images, 
                     selected_recommendations=None, use_humidity=False, logger_screenshots=None,
                     image_orientations=None):
    """Создание приложений. logger_screenshots: список [(номер_логгера, путь_к_скриншоту), ...]"""
    template_value = selected_template
    if hasattr(selected_template, 'get'):
        template_value = selected_template.get()
    # Нормализация типа отчёта
    is_object = template_value in ("Объект хранения", "ОБЪЕКТ ХРАНЕНИЯ ЛЕКАРСТВЕННЫХ СРЕДСТВ")
    is_zone = template_value in ("Зона хранения", "ЗОНА ХРАНЕНИЯ ЛЕКАРСТВЕННЫХ СРЕДСТВ")
    is_fridge = template_value in ("Холодильник/Морозильник", "ХОЛОДИЛЬНИК(БЕЗ ОТКРЫТИЯ)")
    landscape_for_app1_3 = is_object or is_zone  # Альбомная для объект/зона всегда
    logger_screenshots = logger_screenshots or []
    image_orientations = image_orientations or {}
    
    # Функция для определения ориентации страницы на основе ориентации изображения
    def get_page_orientation(image_key, default_orientation='portrait'):
        """Определяет ориентацию страницы на основе ориентации изображения"""
        if image_key in image_orientations:
            return image_orientations[image_key]
        return default_orientation

    # Создание нового стиля для заголовков приложений (если не существует)
    if 'Appendix Heading' not in doc.styles:
        appendix_style = doc.styles.add_style('Appendix Heading', WD_STYLE_TYPE.PARAGRAPH)
        appendix_style.base_style = doc.styles['Heading 1']
        appendix_style.font.name = 'Times New Roman'
        appendix_style.font.size = Pt(12)
        appendix_style.font.bold = True
        appendix_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        appendix_style.paragraph_format.space_before = Pt(0)
        appendix_style.paragraph_format.space_after = Pt(0)
        appendix_style.paragraph_format.left_indent = Inches(0)  # Убираем отступ слева
        appendix_style.paragraph_format.right_indent = Inches(0)  # Убираем отступ справа

    # Функция для добавления заголовка приложения
    def add_appendix_heading(text):
        p = doc.add_paragraph(text, style='Appendix Heading')
        return p

    # Приложение 1: определяем ориентацию на основе изображения планировки
    app1_orientation = get_page_orientation('layout', 'portrait')
    if app1_orientation == 'landscape':
        sect = doc.add_section(WD_SECTION.NEW_PAGE)
        sect.orientation = WD_ORIENT.LANDSCAPE
        sect.page_width = Mm(297)
        sect.page_height = Mm(210)
        sect.top_margin = Mm(20)
        sect.bottom_margin = Mm(20)
        sect.left_margin = Mm(15)
        sect.right_margin = Mm(15)
    else:
        sect = doc.add_section(WD_SECTION.NEW_PAGE)
        sect.orientation = WD_ORIENT.PORTRAIT
        sect.page_width = Mm(210)
        sect.page_height = Mm(297)
        sect.top_margin = Mm(15)
        sect.bottom_margin = Mm(15)
        sect.left_margin = Mm(20)
        sect.right_margin = Mm(20)

    add_appendix_heading("Приложение 1")
    add_appendix_heading("Планировка зоны хранения лекарственных средств")

    # Добавление изображения планировки
    if images['layout']:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)  # Добавляем 6 пт перед изображением
        p.paragraph_format.space_after = Pt(6)   # Добавляем 6 пт после изображения
        run = p.add_run()
        # Определяем ориентацию для изображения планировки
        layout_orientation = image_orientations.get('layout', 'portrait')
        is_landscape = layout_orientation == 'landscape'
        width_inches, height_inches = get_image_size_for_page(images['layout'], is_landscape=is_landscape)
        run.add_picture(images['layout'], width=Inches(width_inches), height=Inches(height_inches))
    else:
        doc.add_paragraph("Изображение планировки не загружено")

    # Добавление разрыва страницы
    doc.add_page_break()

    # Приложение 2: Анализ рисков
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(297)
    section.page_height = Mm(210)
    section.top_margin = Mm(20)  # 2 см сверху
    section.bottom_margin = Mm(20)  # 2 см снизу (для симметрии)
    section.left_margin = Mm(15)  # 1.5 см слева
    section.right_margin = Mm(15)  # 1.5 см справа


    add_appendix_heading("Приложение 2")
    add_appendix_heading("Анализ рисков")
    # Добавляем интервал 6 пт после надписи "Анализ рисков"
    if doc.paragraphs:
        doc.paragraphs[-1].paragraph_format.space_after = Pt(6)



    table = doc.add_table(rows=2, cols=9)
    table.style = 'Table Grid'
    table.autofit = False
    table.allow_autofit = False

    # Установка ширины столбцов
    widths = [0.3, 2, 1, 1, 1, 2, 1, 1, 1]
    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width)

    # Заголовки таблицы
    headers = ["№", "Идентифицированный риск", "Последствия возникнове-ния риска (P)", "Вероятность возникнове-ния риска (S)", 
            "Оценка риска (OR)", "Меры снижения", "Остаточный риск"]
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(9.5)
        cell.width = Inches(widths[i])

    # Объединение ячеек в первых 6 столбцах
    for i in range(6):
        table.cell(0, i).merge(table.cell(1, i))

    # Объединение ячеек для "Остаточный риск" по горизонтали
    table.cell(0, 6).merge(table.cell(0, 8))

    # Добавление подзаголовков для "Остаточный риск"
    table.cell(1, 6).text = "Последствия возникнове-ния риска (P)"
    table.cell(1, 7).text = "Вероятность возникнове-ния риска (S)"
    table.cell(1, 8).text = "Оценка риска (OR)"
    for i in range(6, 9):
        cell = table.cell(1, i)
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(9.5)

    # Добавление данных в таблицу
    for i, risk in enumerate(saved_risk_areas, start=1):
        row = table.add_row()
        row.cells[0].text = str(i)
        if is_fridge:
            row.cells[1].text = f"Риск выхода климатических условий (температуры) за установленные пределы в зоне хранения лекарственных средств {risk}"
        elif is_object or is_zone:
            row.cells[1].text = f"Риск выхода климатических условий (температуры, влажности) за установленные пределы в зоне хранения лекарственных средств {risk}"
        else:
            # По умолчанию для других типов
            row.cells[1].text = f"Риск выхода климатических условий (температуры, влажности) за установленные пределы в зоне хранения лекарственных средств {risk}"
        row.cells[2].text = "3"
        row.cells[3].text = "3"
        row.cells[4].text = "9"
        row.cells[5].text = f"Провести анализ зоны хранения лекарственных средств {risk} и на основании этого анализа обозначить места установки логгеров на схеме размещения"
        row.cells[6].text = "2"
        row.cells[7].text = "1"
        row.cells[8].text = "2"

        # Центрирование текста и выделение жирным шрифтом в ячейках
        for j, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(9.5)
            if j not in [1, 5]:  # Исключаем ячейки с индексами 1 и 5
                for run in paragraph.runs:
                    run.bold = True

    # После заполнения таблицы данными, еще раз устанавливаем ширину
    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width)

    # Принудительно применяем ширину к каждой ячейке
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(widths[i])

    # Добавление разрыва страницы
    doc.add_page_break()


    # Вставка содержимого из rrr.docx сразу после таблицы рисков
    rrr_path = resource_path('rrr.docx')
    if os.path.exists(rrr_path):
        rrr_doc = Document(rrr_path)
    
        # Находим индекс последнего элемента таблицы рисков
        insert_index = -1
        for i, element in enumerate(doc.element.body):
            if isinstance(element, docx.oxml.table.CT_Tbl):
                insert_index = i + 1
    
        # Добавляем разрыв страницы после последней таблицы
        if insert_index != -1:
            page_break = OxmlElement('w:p')
            run = OxmlElement('w:r')
            br = OxmlElement('w:br')
            br.set(qn('w:type'), 'page')
            run.append(br)
            page_break.append(run)
            doc.element.body.insert(insert_index, page_break)
            insert_index += 1
    
        # Копируем содержимое из rrr.docx
        for i, element in enumerate(rrr_doc.element.body):
            new_element = deepcopy(element)
            doc.element.body.insert(insert_index + i, new_element)
    
        # Изменяем интервал для всех параграфов в новой секции
        for paragraph in doc.paragraphs[insert_index:insert_index+len(rrr_doc.paragraphs)]:
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
    
    else:
        print(f"Файл {rrr_path} не найден")

    # Приложение 3: определяем ориентацию на основе изображения схемы размещения логгеров
    app3_orientation = get_page_orientation('loggers', 'portrait')
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    if app3_orientation == 'landscape':
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Mm(297)
        section.page_height = Mm(210)
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)
        section.left_margin = Mm(15)
        section.right_margin = Mm(15)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(15)
        section.bottom_margin = Mm(15)
        section.left_margin = Mm(20)
        section.right_margin = Mm(20)

    add_appendix_heading("Приложение 3")
    add_appendix_heading("Схема размещения логгеров")



    # Добавление изображения планировки
    if images['loggers']:
        # Создаем параграф для изображения и центрируем его
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)  # Добавляем 6 пт перед изображением
        p.paragraph_format.space_after = Pt(6)   # Добавляем 6 пт после изображения
        run = p.add_run()
        # Определяем ориентацию для изображения схемы размещения логгеров
        loggers_orientation = image_orientations.get('loggers', 'portrait')
        is_landscape = loggers_orientation == 'landscape'
        width_inches, height_inches = get_image_size_for_page(images['loggers'], is_landscape=is_landscape)
        run.add_picture(images['loggers'], width=Inches(width_inches), height=Inches(height_inches))
    else:
        doc.add_paragraph("Изображение планировки не загружено")

    # Добавление разрыва страницы
    doc.add_page_break()

    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(297)
    section.page_height = Mm(210)
    section.top_margin = Mm(20)  # 2 см сверху
    section.bottom_margin = Mm(20)  # 2 см снизу (для симметрии)
    section.left_margin = Mm(15)  # 1.5 см слева
    section.right_margin = Mm(15)  # 1.5 см справа

    add_appendix_heading("Приложение 4")
    if is_fridge:
        add_appendix_heading("Графики распределения температуры при проведении исследовании")
        # Добавление заголовка таблицы
        table_heading = doc.add_paragraph()
        run = table_heading.add_run()
        run.add_tab()
        run = table_heading.add_run("На рисунке 4.1 представлен график распределения температуры в зоне хранения лекарственных средств на протяжении всего времени ис-следования. Рисунок 4.1 – График распределения температуры в зоне хранения лекарственных средств на протяжении всего времени исследования")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        table_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        table_heading.paragraph_format.space_after = Pt(6)

        # Добавляем "Добавить график" между блоками
        add_graph_text = doc.add_paragraph()
        add_graph_text.add_run().add_tab()
        run = add_graph_text.add_run("Добавить график")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        add_graph_text.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_graph_text.paragraph_format.space_after = Pt(6)

        # График температуры для первого периода
        if period_images and 1 in period_images and 'temp_fridge' in period_images[1]:
            # Создаем параграф для изображения и центрируем его
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            width_inches, height_inches = get_image_size_for_page(period_images[1]['temp_fridge'], is_landscape=True)
            run.add_picture(period_images[1]['temp_fridge'], width=Inches(width_inches), height=Inches(height_inches))

        # Холодильник: листы 4.2 и 4.3
        doc.add_page_break()
        table_heading = doc.add_paragraph()
        run = table_heading.add_run()
        run.add_tab()
        run = table_heading.add_run("На рисунке 4.2 представлен график распределения температуры около зоны хранения лекарственных средств на протяжении всего времени исследования. Рисунок 4.2 – График распределения температуры около зоны хранения лекарственных средств на протяжении всего времени исследования")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        table_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        table_heading.paragraph_format.space_after = Pt(6)

        # Добавляем "Добавить график" между блоками
        add_graph_text = doc.add_paragraph()
        add_graph_text.add_run().add_tab()
        run = add_graph_text.add_run("Добавить график")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        add_graph_text.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_graph_text.paragraph_format.space_after = Pt(6)

        if period_images and 1 in period_images and 'temp_external' in period_images[1]:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            width_inches, height_inches = get_image_size_for_page(period_images[1]['temp_external'], is_landscape=True)
            run.add_picture(period_images[1]['temp_external'], width=Inches(width_inches), height=Inches(height_inches))
        else:
            doc.add_paragraph("График не загружен")

        doc.add_page_break()
        table_heading = doc.add_paragraph()
        run = table_heading.add_run()
        run.add_tab()
        run = table_heading.add_run("На рисунке 4.3 представлен график распределения относительной влажности около зоны хранения лекарственных средств на протяжении всего времени исследования. Рисунок 4.3 – График распределения относительной влажности около зоны хранения лекарственных средств на протяжении всего времени исследования")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        table_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        table_heading.paragraph_format.space_after = Pt(6)

        if period_images and 1 in period_images and 'humidity_external' in period_images[1]:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            width_inches, height_inches = get_image_size_for_page(period_images[1]['humidity_external'], is_landscape=True)
            run.add_picture(period_images[1]['humidity_external'], width=Inches(width_inches), height=Inches(height_inches))
        else:
            doc.add_paragraph("График не загружен")
    elif is_object or is_zone:
        if use_humidity:
            add_appendix_heading("Графики распределения температуры и влажности при проведении исследований")
        else:
            add_appendix_heading("График распределения температуры при проведении исследований")

            # Добавление заголовка таблицы
            table_heading = doc.add_paragraph()
            run = table_heading.add_run()
            run.add_tab()  # Добавляем табуляцию перед текстом
            run = table_heading.add_run("На рисунке 4.1 представлен график распределения температуры в зоне хранения лекарственных средств на протяжении всего времени ис-следования. Рисунок 4.1 – График распределения температуры в зоне хранения лекарственных средств на протяжении всего времени исследования")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            table_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT        
            table_heading.paragraph_format.space_after = Pt(6)

            # Добавляем "Добавить график" между блоками
            add_graph_text = doc.add_paragraph()
            add_graph_text.add_run().add_tab()
            run = add_graph_text.add_run("Добавить график")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            add_graph_text.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_graph_text.paragraph_format.space_after = Pt(6)

            p_graf = doc.add_paragraph("График температуры")
            for r in p_graf.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(12)

            # График температуры для первого периода
            if period_images and 1 in period_images and 'temp_loggers' in period_images[1]:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                width_inches, height_inches = get_image_size_for_page(period_images[1]['temp_loggers'], is_landscape=True)
                run.add_picture(period_images[1]['temp_loggers'], width=Inches(width_inches), height=Inches(height_inches))
                
                # Добавляем подпись к рисунку
                caption = doc.add_paragraph("Рисунок 4.1 – График распределения температуры в зоне хранения лекарственных средств на протяжении всего времени исследования")
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in caption.runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(12)
                caption.paragraph_format.space_before = Pt(6)

            # Добавление разрыва страницы
            doc.add_page_break()

            # Добавляем влажность только если она учитывается
            if use_humidity:
                # Добавление заголовка таблицы
                table_heading = doc.add_paragraph()
                run = table_heading.add_run()
                run.add_tab()  # Добавляем табуляцию перед текстом
                run = table_heading.add_run("На рисунке 4.2 представлен график распределения относительной влажности в зоне хранения лекарственных средств на протяжении всего времени исследования. Рисунок 4.2 – График распределения относительной влажности в зоне хранения лекарственных средств на протяжении всего времени исследова-ния")
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                table_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT        
                table_heading.paragraph_format.space_after = Pt(6)

                p_graf = doc.add_paragraph("График относительной влажности")
                for r in p_graf.runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(12)

                if period_images and 1 in period_images and 'humidity_loggers' in period_images[1]:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    width_inches, height_inches = get_image_size_for_page(period_images[1]['humidity_loggers'], is_landscape=True)
                    run.add_picture(period_images[1]['humidity_loggers'], width=Inches(width_inches), height=Inches(height_inches))
                    
                    # Добавляем подпись к рисунку
                    caption = doc.add_paragraph("Рисунок 4.2 – График распределения относительной влажности в зоне хранения лекарственных средств на протяжении всего времени исследования")
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in caption.runs:
                        r.font.name = 'Times New Roman'
                        r.font.size = Pt(12)
                    caption.paragraph_format.space_before = Pt(6)

            # Для зоны хранения добавляем дополнительные графики (температура и влажность около зоны)
            if is_zone:
                doc.add_page_break()

                table_heading = doc.add_paragraph()
                run = table_heading.add_run()
                run.add_tab()
                run = table_heading.add_run("На рисунке 4.2 представлен график распределения температуры около зоны хранения лекарственных средств на протяжении всего времени исследования. Рисунок 4.2 – График распределения температуры около зоны хранения лекарственных средств на протяжении всего времени исследования")
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                table_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT        
                table_heading.paragraph_format.space_after = Pt(6)

                # Добавляем "Добавить график" между блоками
                add_graph_text = doc.add_paragraph()
                add_graph_text.add_run().add_tab()
                run = add_graph_text.add_run("Добавить график")
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                add_graph_text.alignment = WD_ALIGN_PARAGRAPH.LEFT
                add_graph_text.paragraph_format.space_after = Pt(6)

                p_graf = doc.add_paragraph("График")
                for r in p_graf.runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(12)

                if period_images and 1 in period_images and 'temp_external' in period_images[1]:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    width_inches, height_inches = get_image_size_for_page(period_images[1]['temp_external'], is_landscape=True)
                    run.add_picture(period_images[1]['temp_external'], width=Inches(width_inches), height=Inches(height_inches))

                # Добавление разрыва страницы
                doc.add_page_break()

                # Добавляем влажность около зоны только если она учитывается
                if use_humidity:
                    # Добавление заголовка таблицы
                    table_heading = doc.add_paragraph()
                    run = table_heading.add_run()
                    run.add_tab()  # Добавляем табуляцию перед текстом
                    run = table_heading.add_run("На рисунке 4.3 представлен график распределения относительной влажности в зоне хранения лекарственных средств на протяжении всего времени исследования. Рисунок 4.3 – График распределения относительной влажности в зоне хранения лекарственных средств на протяжении всего времени исследова-ния")
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    table_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT        
                    table_heading.paragraph_format.space_after = Pt(6)   # Добавляем 6 пт после изображения

                    if period_images and 1 in period_images and 'humidity_loggers' in period_images[1]:
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run()
                        width_inches, height_inches = get_image_size_for_page(period_images[1]['humidity_loggers'], is_landscape=True)
                        run.add_picture(period_images[1]['humidity_loggers'], width=Inches(width_inches), height=Inches(height_inches))

                    # Добавление подписи к рисунку
                    p_graf = doc.add_paragraph("График")
                    for r in p_graf.runs:
                        r.font.name = 'Times New Roman'
                        r.font.size = Pt(12)
                    
                    # Добавляем "Добавить график" между блоками
                    add_graph_text = doc.add_paragraph()
                    add_graph_text.add_run().add_tab()
                    run = add_graph_text.add_run("Добавить график")
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    add_graph_text.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    add_graph_text.paragraph_format.space_after = Pt(6)

                    # Добавление разрыва страницы
                    doc.add_page_break()

                    # Добавление заголовка таблицы
                    table_heading = doc.add_paragraph()
                    run = table_heading.add_run()
                    run.add_tab()  # Добавляем табуляцию перед текстом
                    run = table_heading.add_run("На рисунке 4.4 представлен график распределения относительной влажности около зоны хранения лекарственных средств на протяжении всего времени исследования. Рисунок 4.4 – График распределения относительной влажности около зоны хранения лекарственных средств на протяжении всего времени исследования")
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    table_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT       
                    table_heading.paragraph_format.space_after = Pt(6)

                    p_graf = doc.add_paragraph("График")
                    for r in p_graf.runs:
                        r.font.name = 'Times New Roman'
                        r.font.size = Pt(12)

                    if period_images and 1 in period_images and 'humidity_external' in period_images[1]:
                        # Создаем параграф для изображения и центрируем humidity_external
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run()
                        width_inches, height_inches = get_image_size_for_page(period_images[1]['humidity_external'], is_landscape=True)
                        run.add_picture(period_images[1]['humidity_external'], width=Inches(width_inches), height=Inches(height_inches))
                        
                        # Добавляем подпись к рисунку
                        caption = doc.add_paragraph("Рисунок 4.4 – График распределения относительной влажности около зоны хранения лекарственных средств на протяжении всего времени исследования")
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for r in caption.runs:
                            r.font.name = 'Times New Roman'
                            r.font.size = Pt(12)
                        caption.paragraph_format.space_before = Pt(6)

    # Приложение 5: Графики по логгерам (динамические)
    if logger_screenshots:
        section5 = doc.add_section(WD_SECTION.NEW_PAGE)
        if is_object or is_zone:
            section5.orientation = WD_ORIENT.LANDSCAPE
            section5.page_width = Mm(297)
            section5.page_height = Mm(210)
            section5.top_margin = Mm(20)
            section5.bottom_margin = Mm(20)
            section5.left_margin = Mm(15)
            section5.right_margin = Mm(15)
        else:
            section5.orientation = WD_ORIENT.PORTRAIT
            section5.page_width = Mm(210)
            section5.page_height = Mm(297)
            section5.top_margin = Mm(15)
            section5.bottom_margin = Mm(15)
            section5.left_margin = Mm(20)
            section5.right_margin = Mm(20)
        add_appendix_heading("Приложение 5")
        add_appendix_heading("Графики распределения температуры и влажности по каждому логгеру при проведении исследований")
        add_appendix_heading("(скриншоты из ПО Eclerk-2.0)")
        for idx, (logger_num, img_path) in enumerate(logger_screenshots, 1):
            p_intro = doc.add_paragraph()
            run = p_intro.add_run()
            run.add_tab()
            run = p_intro.add_run(f"На рисунке 5.{idx} представлен график распределения температуры и относительной влажности по данным логгера №{logger_num} за весь период исследования.")
            for r in p_intro.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(12)
            if img_path and os.path.exists(img_path):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(6)  # Добавляем 6 пт перед изображением
                p.paragraph_format.space_after = Pt(6)   # Добавляем 6 пт после изображения
                run = p.add_run()
                width_inches, height_inches = get_fixed_image_size_for_logger_screenshots()
                run.add_picture(img_path, width=Inches(width_inches), height=Inches(height_inches))
            else:
                doc.add_paragraph("Изображение не загружено")
            caption = doc.add_paragraph(
                f"Рисунок 5.{idx} – График распределения температуры и относительной влажности по данным логгера №{logger_num} за весь период исследования"
            )
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in caption.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(12)
            caption.paragraph_format.space_before = Pt(6)
            if idx < len(logger_screenshots):
                doc.add_page_break()

    # Приложение 6: Температурная и влажностная карты
    section6 = doc.add_section(WD_SECTION.NEW_PAGE)
    
    # Определяем ориентацию для температурной карты
    if is_fridge:
        temp_orientation = 'portrait'  # Для холодильника всегда книжная
    else:
        # Для объекта/зоны используем ориентацию температурной карты
        temp_orientation = get_page_orientation('temp_map', 'landscape')
    
    if temp_orientation == 'landscape':
        section6.orientation = WD_ORIENT.LANDSCAPE
        section6.page_width = Mm(297)
        section6.page_height = Mm(210)
        section6.top_margin = Mm(20)
        section6.bottom_margin = Mm(20)
        section6.left_margin = Mm(15)
        section6.right_margin = Mm(15)
    else:
        section6.orientation = WD_ORIENT.PORTRAIT
        section6.page_width = Mm(210)
        section6.page_height = Mm(297)
        section6.top_margin = Mm(15)
        section6.bottom_margin = Mm(15)
        section6.left_margin = Mm(20)
        section6.right_margin = Mm(20)
    
    add_appendix_heading("Приложение 6")
    if is_fridge:
        add_appendix_heading("Температурная карта")
    elif is_object or is_zone:
        if use_humidity:
            add_appendix_heading("Температурная и влажностная карты")
        else:
            add_appendix_heading("Температурная карта")

    if is_fridge:
        table_heading = doc.add_paragraph()
        run = table_heading.add_run()
        run.add_tab()
        run = table_heading.add_run("На рисунке 6.1 представлена температурная карта зоны хранения лекарственных средств за весь период исследования.")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        table_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if images.get('temp_map'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            # Определяем ориентацию для изображения температурной карты
            temp_map_orientation = image_orientations.get('temp_map', 'portrait')
            is_landscape = temp_map_orientation == 'landscape'
            width_inches, height_inches = get_image_size_for_page(images['temp_map'], is_landscape=is_landscape)
            run.add_picture(images['temp_map'], width=Inches(width_inches), height=Inches(height_inches))
        else:
            doc.add_paragraph("Изображение планировки не загружено")
        caption = doc.add_paragraph("Рисунок 6.1 – Температурная карта зоны хранения лекарственных средств за весь период исследования")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in caption.runs:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(12)
    elif is_object or is_zone:
        table_heading = doc.add_paragraph()
        run = table_heading.add_run()
        run.add_tab()
        run = table_heading.add_run("На рисунке 6.1 представлена температурная карта зоны хранения лекарственных средств за весь период исследования.")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        table_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if images.get('temp_map'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run()
            # Определяем ориентацию для изображения температурной карты
            temp_map_orientation = image_orientations.get('temp_map', 'portrait')
            is_landscape = temp_map_orientation == 'landscape'
            width_inches, height_inches = get_image_size_for_page(images['temp_map'], is_landscape=is_landscape)
            run.add_picture(images['temp_map'], width=Inches(width_inches), height=Inches(height_inches))
        else:
            doc.add_paragraph("Изображение планировки не загружено")
        caption = doc.add_paragraph("Рисунок 6.1 – Температурная карта зоны хранения лекарственных средств за весь период исследования")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in caption.runs:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(12)

        # Добавляем влажностную карту только если она учитывается
        if use_humidity:
            # Создаем новую секцию для влажностной карты (без разрыва страницы)
            section6_humidity = doc.add_section(WD_SECTION.CONTINUOUS)
            
            # Определяем ориентацию для влажностной карты
            humidity_orientation = get_page_orientation('humidity_map', 'portrait')
            
            if humidity_orientation == 'landscape':
                section6_humidity.orientation = WD_ORIENT.LANDSCAPE
                section6_humidity.page_width = Mm(297)
                section6_humidity.page_height = Mm(210)
                section6_humidity.top_margin = Mm(20)
                section6_humidity.bottom_margin = Mm(20)
                section6_humidity.left_margin = Mm(15)
                section6_humidity.right_margin = Mm(15)
            else:
                section6_humidity.orientation = WD_ORIENT.PORTRAIT
                section6_humidity.page_width = Mm(210)
                section6_humidity.page_height = Mm(297)
                section6_humidity.top_margin = Mm(15)
                section6_humidity.bottom_margin = Mm(15)
                section6_humidity.left_margin = Mm(20)
                section6_humidity.right_margin = Mm(20)

            table_heading = doc.add_paragraph()
            run = table_heading.add_run()
            run.add_tab()
            run = table_heading.add_run("На рисунке 6.2 представлена влажностная карта зоны хранения лекарственных средств за весь период исследования.")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            table_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if images.get('humidity_map'):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run()
                # Определяем ориентацию для изображения влажностной карты
                humidity_map_orientation = image_orientations.get('humidity_map', 'portrait')
                is_landscape = humidity_map_orientation == 'landscape'
                width_inches, height_inches = get_image_size_for_page(images['humidity_map'], is_landscape=is_landscape)
                run.add_picture(images['humidity_map'], width=Inches(width_inches), height=Inches(height_inches))
            else:
                doc.add_paragraph("Изображение планировки не загружено")
            caption = doc.add_paragraph("Рисунок 6.2 – Влажностная карта зоны хранения лекарственных средств за весь период исследования")
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in caption.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(12)
